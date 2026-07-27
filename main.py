import os
import re
import json
from datetime import datetime
from io import BytesIO
from typing import Optional, List

# ─────────────────────────────────────────────────────────────
# BLUE JEANS NOVEL ENGINE — main.py 변경 이력 (누적)
#
# v3.15.0 (2026-07-26) — 설계 전달 경로 복구 + 연속성 + 후처리 정리기
#   [실증 기반] 《사랑한다고 했잖아》 UNIT 01·02 원고와 저장 JSON 3종
#   (Unit설계_20260725_0255 / novelengine_20260725_1313)으로 사고 5종 확인.
#
#   1. M19 설계안 전달 — Stage A/B/C 호출부에 all_blueprints_text 전달.
#      v3.14까지 Chapter 1은 Unit 설계안을 한 글자도 받지 못했다.
#      그 결과 설계 "택시 안에서 도균과 조우"가 원고에서 "정류장 조우"로
#      바뀌고 Plant 오브제(택시·칼가방)가 소멸했다.
#   2. sanitize_manuscript() 신설 — 본문 중간 '끝.'·마크다운 마커 제거.
#      Stage 생성 직후 / 병합 직후 / Unit 저장 직후 3중 적용.
#      ★ 문장은 한 글자도 바꾸지 않는다 (작업 규칙 7 준수). ★
#   3. export_docx 제목 오인 수정 — '첫 줄이고 80자 미만'이면 무조건
#      작품 제목으로 렌더해서, [CHAPTER n] 헤더가 빠진 Unit에서
#      첫 대사가 16pt 굵게 중앙정렬로 박혀 나갔다. (UNIT 02 실제 사고)
#      _looks_like_doc_title()로 실제 제목일 때만 제목 처리.
#   4. 상태 메시지 sticky 해소 — '결과가 중간에서 끊겼습니다' 오류가
#      다음 set_status()까지 화면 상단에 영구히 남았다. 1회 표시 후 소거.
#   5. Chapter 1 병합 경로 품질 경고 연결 — v3.14까지 리포트만 조용히
#      저장해 '있었다' 16회(임계치 10)가 그대로 통과했다.
#   6. AUTO_REGEN_MAX_RETRIES 1 → 2.
#   7. M17 Continuity Ledger — generate_continuity_ledger() /
#      gather_continuity_ledger(). 착의·소지품·위치·시각·날씨·인지 원장을
#      Unit 확정 시 추출해 다음 Unit 프롬프트에 주입. STEP 5에서 작가가
#      직접 수정 가능. session_state 키 'continuity_ledger' 신규.
#   8. 진단 지표 4종 추가 — 신체반응 어휘(임계 6) / 서술자 라벨링(1) /
#      회상 마커(3) / 잔여 마커(0). 실측 검증: UNIT 01 신체반응 11회·
#      회상 마커 4회·잔여 마커 2개, UNIT 02 서술자 라벨링 2회 검출.
#   9. STEP 5 도구 3종 신설 — [이 Unit 재검사] [본문 정리] [상태 원장 생성].
#
# v3.15.1 (2026-07-26) — 챕터 제목 폴백 + 교차 Unit 반복 진단 + 통합본 정리
#   1. 챕터 제목 폴백 — 모델이 [CHAPTER n] 헤더를 빼먹으면 chapter_titles가
#      빈 값으로 남는다. extract_blueprint_chapter_title()로 설계안의
#      확정 제목을 채운다. Chapter 1 병합 경로와 Unit 02~12 경로 양쪽 적용.
#      실측: 실제 설계 파일에서 UNIT 01~10 제목 10개 전부 추출 성공.
#   2. analyze_cross_unit_repetition() 신설 — analyze_unit_quality()는 한 Unit
#      내부만 본다. 실제 사고는 Unit을 건너뛰며 났다.
#      "심장이 늑골을 두드렸다"(01) ↔ "심장이 갈비뼈를 쳤다"(02) — 같은 말이다.
#      동의어군 6종 + 특징 어구(6~14자) 재사용을 직전 3개 Unit과 대조.
#      실측: UNIT 02 검사에서 손끝(7→3회)·심장(2→2회)·겹치는 어구 6건 검출.
#      사고 패턴 A 대응 — 정수 키·문자열 키 양쪽 조회.
#   3. final_manuscript_text()에 sanitize_manuscript() 적용 — 옛 Unit 본문의
#      중간 '끝.'이 STEP 7 통합본에 그대로 실려 나가던 문제. 마지막 '끝.'은 보존.
#   ★ 세 항목 모두 진단·도구 제공에 그치고 본문 문장은 수정하지 않는다. ★
#
# v3.15.2 (2026-07-26) — UI 동선 정리 (엔진 룰 변경 없음)
#   1. STEP 6(가제 검토) ↔ STEP 7(저장) 순서 교체 + 접힘 전환.
#      작품 전체 제목은 완고 후 1회짜리 작업인데, Unit을 한 편 생성할 때마다
#      저장하러 내려가는 길목에 STEP 6으로 끼어 있었다.
#   2. 전 STEP 버튼에 실행 순번 부여.
#      생성 순서 = 숫자(5-1~5-4), 사후 점검 도구 = 문자(점검 A·B·C).
#      '해야 하는 순서'와 '필요할 때 쓰는 도구'를 시각적으로 갈랐다.
#   3. Chapter 1 4단계 진행 표시(✅/⬜) 추가.
#   4. STEP 2·3·4·5·6 헤더에 실행 순서 한 줄 안내 추가.
#
# v3.16.0 (2026-07-27) — 회차 제목 단일화 (설계안 = 정본)
#   [배경] 제목이 STEP 4 설계와 STEP 5 집필 두 군데서 따로 만들어졌다.
#   둘이 연결돼 있지 않아 어긋났고, 모델이 집필 단계 지시를 무시해서
#   UNIT 01·02 원고에는 [CHAPTER n] 줄이 아예 없었다. 그 빈 제목이
#   v3.15에서 고친 DOCX 제목 오인 사고의 근본 원인이었다.
#
#   1. resolve_chapter_title() 신설 — 설계안 제목이 정본.
#      구버전 원고가 다른 서브타이틀을 갖고 있으면 설계안이 이기되 알린다.
#      Chapter 1 병합 / Unit 02~12 / 에필로그 3개 경로 모두 적용.
#   2. gather_chapter_titles_text() 신설 — 현재 목차 수집.
#   3. STEP 7을 탭 2개로 재구성 — [작품 제목] [회차 제목 검수].
#      7-2는 새로 짓지 않고 본문 이탈·목차 스포일러·이미지 중복만 검수.
#      제목 교체는 '회차 제목 직접 수정' 폼에서 작가가 직접 한다(규칙 7).
#   4. session_state 키 'chapter_title_review' 신규.
#
# v3.16.1 (2026-07-27) — M20 회차 제목 유형 체계 (prompt.py 주도)
#   설계안 제목 10개가 전부 '수식어 + 추상 명사'라 목차가 한 덩어리였다.
#   제목 유형을 비트 기능에 종속시켜(인물 소개 회차 = 인물 이름,
#   전환점 회차 = 사건 이름) 목차에서 회차가 구별되게 했다.
#   main.py 변경은 없다. 기존 저장 설계안과의 호환은 회귀 테스트로 확인.
#
# v3.16.2 (2026-07-27) — DOCX 빈 줄 처리 + 집필 메타 발화 검출
#   1. export_docx(title, content, spacing_mode) — 세 번째 인자 신규.
#      기존에는 본문의 빈 줄 하나하나를 장면 전환으로 보고 빈 단락을 넣었다.
#      모델은 모든 문단 사이에 빈 줄을 넣으므로 절반이 빈 단락이 됐다.
#      실측 — UNIT 01: 219단락 중 108개(49%)가 빈 단락.
#      standard(기본): 빈 줄 1개는 문단 구분으로만 처리(빈 단락 없음).
#      relaxed: standard + 문단 뒤 6pt 여백.
#      web: 구버전 동작 유지.
#      빈 줄 2개 이상 / 장면 전환 마커(*, ·, — 등)는 모든 모드에서 보존.
#      실측 결과 — 220단락 → 112단락, 빈 단락 108 → 0, 장면 전환은 유지.
#   2. STEP 6에 'DOCX 문단 간격' 라디오 추가.
#   3. sanitize_manuscript에 집필 메타 발화 검출 추가.
#      "Stage B가 닫혔으니 ~ Stage C를 열겠습니다"가 본문에 섞인 사고.
#      ★ 본문 문장이므로 자동 삭제하지 않고 경고만 한다(규칙 7). ★
#
# v3.16.3 (2026-07-27) — master 서식 통일 + 원장 잘림 수정 + 회차 재생성
#   1. export_docx 서식을 작가 master 파일 규격으로 통일.
#      여백 상하좌우 2.0cm / 함초롱바탕 10.5pt / 줄간격 1.15 / 양쪽정렬.
#      (기존: 상3·하좌우2.54cm, 바탕 10pt, 줄간격 1.6)
#      작가가 master에 붙여넣을 때 서식이 흔들리지 않게 하는 것이 목적.
#   2. generate_continuity_ledger max_tokens 1500 → 4000.
#      실제 사고 — UNIT 01 원장이 [착의] 항목 중간에서 잘려 뒤쪽 6개 항목이
#      사라졌고, "명함을 택시 안에서 받았다"가 원장에 없어 UNIT 02가
#      "파티 케이터링 화장실 앞"으로 왜곡해도 걸러지지 않았다.
#      필수 항목 누락·잘림 검사 추가. 실패 시 조용히 넘어가지 않고 경고.
#   3. sanitize_manuscript에 기울임 마크다운 제거 추가.
#      실제 사고 — 문자 메시지가 *오늘도 늦어? 밥은 먹었고.* 로 나갔다.
#      장면 전환 마커(별표만 있는 줄)와 구분해 내용이 있을 때만 벗긴다.
#   4. 회차 재생성 UI 신설.
#      · [♻️ 1화 처음부터 다시 쓰기] — Stage A·B·C + UNIT 01 확정본 +
#        제목 + 원장 + 요약을 한 번에 비운다. 설계안은 보존.
#      · [♻️ N화 처음부터 다시 쓰기] — 선택 Unit을 비운다.
#      둘 다 체크박스 확인을 거쳐야 실행된다.
#      기존 '다시 쓰기'는 원고를 손보는 것이고, 이것은 새로 뽑는 것이다.
#      비우지 않으면 옛 원고가 previous_drafts로 섞여 들어간다.
#
# © 2026 BLUE JEANS PICTURES. All rights reserved.
# ─────────────────────────────────────────────────────────────

import streamlit as st

try:
    import anthropic
except ImportError:
    anthropic = None

from docx import Document

from prompt import (
    SYSTEM_PROMPT,
    STYLE_DNA_ANALYSIS_PROMPT,
    build_locked_block,
    build_merge_analysis_prompt,
    build_gap_diagnosis_prompt,
    build_story_reinforcement_prompt,
    build_unit_blueprint_prompt,
    build_unit_draft_prompt,
    build_unit_rewrite_prompt,
    build_title_review_prompt,
    build_chapter_title_review_prompt,
    build_epilogue_prompt,
    build_expand_incomplete_unit_prompt,
    build_ch1_stage_a_prompt,
    build_ch1_stage_b_prompt,
    build_ch1_stage_c_prompt,
    build_metric_watchlist_block,
    SIGNATURE_FOOD_OPENING_BLOCK,
    # v3.15 신규 (M17 / M19)
    extract_unit_blueprint,
    extract_blueprint_chapter_title,
    build_blueprint_adherence_block,
    build_continuity_ledger_prompt,
    build_continuity_block,
    # v3.0 신규
    NOVEL_ENGINE_VERSION,
    NOVEL_ENGINE_BUILD_DATE,
    NOVEL_ENGINE_VERSION_TAG,
    get_novel_engine_version_info,
    _PROFESSION_PACK_AVAILABLE,
    _PERIOD_PACK_AVAILABLE,
)

# v3.0 Period Pack 키 목록 (STEP 1 선택지용)
try:
    from period_pack import get_all_period_keys, get_period_label, detect_period_from_locked
    PERIOD_KEYS = get_all_period_keys()
except ImportError:
    PERIOD_KEYS = []
    def get_period_label(k): return ""
    def detect_period_from_locked(t): return []

# v3.1 시나리오 → 소설화 추출 모듈
try:
    from scenario_extractor import (
        extract_text_from_docx,
        extract_text_from_txt,
        analyze_scenario_structure,
        extract_scenario_fields,
        build_unit_mapping_text,
        SCENARIO_EXTRACTOR_VERSION,
    )
    _SCENARIO_EXTRACTOR_AVAILABLE = True
except ImportError:
    _SCENARIO_EXTRACTOR_AVAILABLE = False
    def extract_text_from_docx(b): return ""
    def extract_text_from_txt(b): return ""
    def analyze_scenario_structure(t): return {}
    def extract_scenario_fields(t, c, **kw): return {"_error": "scenario_extractor 미로드"}
    def build_unit_mapping_text(m): return ""
    SCENARIO_EXTRACTOR_VERSION = "미로드"

# v3.2 Creator Engine JSON → 소설화 변환 모듈
try:
    from creator_extractor import (
        load_creator_json,
        is_creator_json,
        get_creator_meta,
        extract_creator_fields,
        CREATOR_EXTRACTOR_VERSION,
    )
    _CREATOR_EXTRACTOR_AVAILABLE = True
except ImportError:
    _CREATOR_EXTRACTOR_AVAILABLE = False
    def load_creator_json(b): return {"_error": "creator_extractor 미로드"}
    def is_creator_json(d): return False
    def get_creator_meta(d): return {}
    def extract_creator_fields(d, c, **kw): return {"_error": "creator_extractor 미로드"}
    CREATOR_EXTRACTOR_VERSION = "미로드"

# v3.3 Idea Engine JSON → 소설화 변환 모듈
try:
    from idea_extractor import (
        load_idea_json,
        is_idea_json,
        get_idea_meta,
        collect_pending_items,
        extract_idea_fields,
        IDEA_EXTRACTOR_VERSION,
    )
    _IDEA_EXTRACTOR_AVAILABLE = True
except ImportError:
    _IDEA_EXTRACTOR_AVAILABLE = False
    def load_idea_json(b): return {"_error": "idea_extractor 미로드"}
    def is_idea_json(d): return False
    def get_idea_meta(d): return {}
    def collect_pending_items(d): return []
    def extract_idea_fields(d, c, **kw): return {"_error": "idea_extractor 미로드"}
    IDEA_EXTRACTOR_VERSION = "미로드"

# ─────────────────────────────────────
# CONFIG
# ─────────────────────────────────────
APP_TITLE = "NOVEL ENGINE"
# v3.11 — 엔진 정체성. 만화=그래픽 노블처럼, 이 엔진의 결과물은 시네마틱 노블이다.
APP_FORMAT = "CINEMATIC NOVEL"
APP_FORMAT_DESC = "영화를 연상시키는 속도감 · 3막 15비트 구조 · 소설만의 묘사와 심리"
APP_SUB = "NOVEL WRITER STUDIO v3.0"

DEFAULT_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-5")
MODEL_OPUS = os.getenv("ANTHROPIC_MODEL_OPUS", "claude-opus-4-8")
# v3.4.2 — 토큰 상한 재정비.
# 구모델(8192 상한) 기준으로 잡혀 있던 값들이 신모델(Sonnet 5 / Opus 4.8,
# 최대 출력 128k)에서 오히려 결과를 잘라먹는 원인이 됐다.
# 실제 사고: 입력 자료가 풍부한 작품에서 STEP 2·3·4 결과가 문장 중간에 끊김
# ("기태의 위협 방식(스토킹" 에서 종료되는 식).
MAX_TOKENS_SHORT = 4000       # 짧은 회신 (제목 검토 등)
MAX_TOKENS_MID = 6000         # 중간 길이 (레거시 유지 — 참조하는 곳 남아있음)
MAX_TOKENS_LONG = 16000       # Unit 본문 집필
MAX_TOKENS_ANALYSIS = 16000   # STEP 2 분석·진단 리포트
MAX_TOKENS_DESIGN = 20000     # STEP 3 기승전결 보강, STEP 4 Unit 설계
MAX_TOKENS_EXTRACT = 32000    # STEP 0 JSON 추출 (가장 긴 출력)
# v3.3.1 — JSON 구조 추출 전용 상한.
# Creator/Idea 다이제스트가 3만 자를 넘고 12 Unit 매핑까지 한 번에 뽑아야 해서
# 8192로는 JSON이 중간에 잘릴 수 있다. 신모델(Sonnet 5 / Opus 4.8)은
# 최대 출력 128k를 지원하므로 추출 단계만 상향한다.
# v3.3.3: 16000도 부족한 사례가 있어 32000으로 재상향.
MAX_TOKENS_EXTRACT = 32000

# v3.0 M1: BJND Scene Enforcer 임계치 (사전 차단 + 자동 재생성 트리거)
BJND_THRESHOLDS = {
    "있었다": 10,        # v2.5는 15 → v3.0은 10 (강화)
    "것이었다": 2,       # v2.5는 3 → v3.0은 2 (강화)
    "대사태그": 12,      # 말했다+물었다+대답했다 합계
    "마치처럼": 1,       # "마치~처럼"/"~듯했다"/"~같았다" 합계
    "현재형": 3,         # 현재형 종결 (치명적)
    "계량수치": 2,       # v3.12 M15 — 숫자+계량단위 / 소수점 (서술문)
    # v3.15 신규 진단 지표
    "신체반응": 6,       # 손끝·심장·등골·목덜미 등 신체 반응 어휘 (M11 보강)
    "기억오류": 1,       # "~하지는 않았다" 류 감정 라벨링·기억 오류 장치 (M18)
    "회상마커": 3,       # Chapter 1 회상 진입 마커 (M12 / Stage B 회상 금지)
}

AUTO_REGEN_MAX_RETRIES = 2  # v3.15 — 1회로는 임계치 초과가 통과되는 사례가 있어 2회로 상향

# v3.15 신규 진단 정규식
# ① 신체 반응 어휘 — 같은 부위로만 감정을 번역하는 클리셰 반복 감지.
#    실측: 《사랑한다고 했잖아》 UNIT 01·02에서 "손끝" 10회, 심장 비유 중복.
BODY_REACTION_RE = re.compile(
    r"(손끝|손가락 끝|심장|가슴이 (?:쿵|철렁)|늑골|갈비뼈|등골|등이 (?:서늘|차가)|"
    r"목덜미|뒷목|숨이 (?:막|멎)|어깨가 (?:굳|움츠))"
)
# ② 감정 라벨링 / 기억 오류 장치 — 서술자가 독자에게 정답을 짚어주는 구문.
#    강력한 장치지만 반복되면 기법이 노출된다. Unit당 1회까지.
NARRATOR_LABEL_RE = re.compile(
    r"(?:이라고|라고|으로)?\s*(?:이름 붙이지|명명하지|따져보지|묻지|생각하지|"
    r"헤아리지|의심하지|확인하지)\s*(?:는)?\s*않았다"
)
# ③ Chapter 1 회상 진입 마커 — Stage B는 회상 절대 금지, Chapter 1 전체도 30% 제한.
FLASHBACK_MARKER_RE = re.compile(
    r"(지난(?:달|주|해|번)에도|그때도|어릴 때|어린 시절|처음 만났을 때|"
    r"몇 년 전|넉 달 전|반년 전|작년|재작년|그 무렵|예전에|한때)"
)

# v3.12 M15 / v3.13 M15-B — 계량 수치 감지 정규식
# 본문 검사(analyze_unit_quality)와 자료 스캔(scan_metric_expressions)의
# 기준이 다르다. 자료의 %는 장르 배합 비중 같은 집필 지시이고,
# 순수 소수점은 엔진 버전·단서 강도 표기라서 자료 스캔에서 제외한다.
#
# 단위를 3군으로 나눈다.
#   A군 긴 단위  — 숫자와 단위 사이 공백 허용 (충돌 위험 없음)
#   B군 기호 단위 — 공백 불허 (168cm 형태)
#   C군 1글자 단위 — 공백 불허 + 뒤 조사 화이트리스트
#     ("Unit 10은 … 도균" 에서 '10 도'를 잡는 오탐을 막는다)
_METRIC_A = r"센티미터|센티|밀리미터|밀리|센치|미터|킬로그램|킬로|그램|리터|밀리리터|인치|피트|야드"
_METRIC_B = r"㎝|㎜|cm|mm|㎏|kg|㎖|ml|평"
_METRIC_B_PCT = r"퍼센트|프로|%"
_METRIC_C = r"도|초|박"
_METRIC_C_TAIL = r"(?=$|[^가-힣]|까지|씩|간|가|를|은|는|의|로|만|나|짜리|어긋)"

# 본문용 — % 포함
METRIC_UNIT_RE = (
    rf"\d+(?:\.\d+)?\s*(?:{_METRIC_A})"
    rf"|\d+(?:\.\d+)?(?:{_METRIC_B}|{_METRIC_B_PCT})"
    rf"|\d+(?:\.\d+)?(?:{_METRIC_C}){_METRIC_C_TAIL}"
)
METRIC_DECIMAL_RE = r"\d+\.\d+"

# 자료용 — % 제외, 순수 소수점 제외 (설계 지시·버전 표기 오탐 차단)
METRIC_SCAN_RE = (
    rf"\d+(?:\.\d+)?\s*(?:{_METRIC_A})"
    rf"|\d+(?:\.\d+)?(?:{_METRIC_B})"
    rf"|\d+(?:\.\d+)?(?:{_METRIC_C}){_METRIC_C_TAIL}"
)

# 자료 스캔에서 건너뛸 Creator JSON 키 (메타데이터·식별자)
METRIC_SCAN_SKIP_KEYS = ("version", "engine", "_meta", "_id", "id", "score", "ratio")

UNIT_TARGET_LENGTHS = {
    1: 7000, 2: 7000, 3: 8000,
    4: 8000, 5: 8000, 6: 9000,
    7: 9000, 8: 9000, 9: 8000,
    10: 8000, 11: 8000, 12: 9000,
    13: 2500,
}

UNIT_MIN_LENGTHS = {
    1: 6000, 2: 6000, 3: 6500,
    4: 6500, 5: 6500, 6: 7000,
    7: 7000, 8: 7000, 9: 6500,
    10: 6500, 11: 6500, 12: 7000,
    13: 1800,
}

# ─────────────────────────────────────
# PAGE
# ─────────────────────────────────────
st.set_page_config(
    page_title="BLUE JEANS | Novel Engine",
    page_icon="👖",
    layout="wide",
)

# ─────────────────────────────────────
# CSS
# ─────────────────────────────────────
st.markdown(
    """
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://cdn.jsdelivr.net/gh/projectnoonnu/2408-3@latest/Paperlogy.css');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&display=swap');

:root {
    --navy: #202A78;
    --y: #FFCB05;
    --bg: #F7F7F5;
    --card: #FFFFFF;
    --card-border: #DDDDE6;
    --t: #2A2A3A;
    --dim: #8A8FA3;
    --light-bg: #EEEEF6;
    --display: 'Playfair Display', 'Paperlogy', 'Georgia', serif;
    --body: 'Pretendard', -apple-system, sans-serif;
    --heading: 'Paperlogy', 'Pretendard', sans-serif;
}

html, body, [class*="css"] {
    font-family: var(--body);
    color: var(--t);
    -webkit-font-smoothing: antialiased;
}
.stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"],
[data-testid="stMainBlockContainer"], [data-testid="stHeader"],
[data-testid="stBottom"] {
    background-color: var(--bg) !important;
    color: var(--t) !important;
}
.stMarkdown, .stText, .stCode { color: var(--t) !important; }
h1,h2,h3,h4,h5,h6 {
    color: var(--navy) !important;
    font-family: var(--heading) !important;
}
p, span, label, div, li { color: var(--t); }
section[data-testid="stSidebar"] { display: none; }

.stTextInput input, .stTextArea textarea,
[data-testid="stTextInput"] input, [data-testid="stTextArea"] textarea {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-size: 0.92rem !important;
    padding: 0.65rem 0.85rem !important;
}
.stTextInput input:focus, .stTextArea textarea:focus,
[data-testid="stTextInput"] input:focus, [data-testid="stTextArea"] textarea:focus {
    border-color: var(--navy) !important;
    box-shadow: 0 0 0 2px rgba(32,42,120,0.08) !important;
}
.stTextInput input::placeholder, .stTextArea textarea::placeholder,
[data-testid="stTextInput"] input::placeholder, [data-testid="stTextArea"] textarea::placeholder {
    color: var(--dim) !important;
    font-size: 0.85rem !important;
}
.stSelectbox > div > div, [data-baseweb="select"] > div, [data-baseweb="select"] input {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border-color: var(--card-border) !important;
    border-radius: 8px !important;
}
[data-baseweb="popover"], [data-baseweb="menu"], [role="listbox"], [role="option"] {
    background-color: var(--card) !important;
    color: var(--t) !important;
}
[role="option"]:hover { background-color: var(--light-bg) !important; }
.stTextInput label, .stTextArea label, .stSelectbox label {
    color: var(--t) !important;
    font-weight: 600 !important;
    font-size: 0.82rem !important;
    margin-bottom: 0.3rem !important;
}
.stButton > button {
    color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important;
    background-color: var(--card) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-weight: 700 !important;
    font-size: 0.88rem !important;
    padding: 0.55rem 1.2rem !important;
    transition: all 0.2s;
}
.stButton > button:hover {
    border-color: var(--navy) !important;
    box-shadow: 0 2px 8px rgba(32,42,120,0.08) !important;
}
.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background-color: var(--y) !important;
    color: var(--navy) !important;
    border-color: var(--y) !important;
    font-weight: 800 !important;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background-color: #E8B800 !important;
    box-shadow: 0 2px 12px rgba(255,203,5,0.3) !important;
}
.stDownloadButton > button {
    color: var(--navy) !important;
    border: 1.5px solid var(--y) !important;
    background-color: var(--y) !important;
    border-radius: 8px !important;
    font-family: var(--body) !important;
    font-weight: 800 !important;
    font-size: 0.88rem !important;
    padding: 0.55rem 1.2rem !important;
}
.stExpander, details, details summary {
    background-color: var(--card) !important;
    color: var(--t) !important;
    border: 1px solid var(--card-border) !important;
    border-radius: 8px !important;
}
details[open] > div { background-color: var(--card) !important; }

.header-wrap {
    text-align: center;
    padding: 2.5rem 0 0.5rem 0;
}
.header {
    font-size: 0.82rem;
    font-weight: 700;
    color: var(--navy);
    letter-spacing: 0.25em;
    font-family: var(--body);
    text-transform: uppercase;
    margin-bottom: 0.6rem;
}
.brand-title {
    font-size: 2.8rem;
    font-weight: 900;
    color: var(--navy);
    font-family: var(--display);
    letter-spacing: -0.01em;
    position: relative;
    display: inline-block;
}
.brand-title::after {
    content: '';
    position: absolute;
    bottom: 2px;
    left: 0;
    width: 100%;
    height: 4px;
    background: var(--y);
    border-radius: 2px;
}
.tagline {
    font-size: 0.92rem;
    font-weight: 700;
    color: var(--navy);
    font-family: var(--display);
    letter-spacing: 0.14em;
    margin-top: 0.7rem;
    text-transform: uppercase;
}
.tagline span {
    background: linear-gradient(transparent 62%, var(--y) 62%);
    padding: 0 0.15em;
}
.tagline-desc {
    font-size: 0.74rem;
    color: var(--dim);
    letter-spacing: 0.02em;
    margin-top: 0.3rem;
    font-family: var(--body);
}
.sub {
    font-size: 0.68rem;
    color: var(--dim);
    letter-spacing: 0.22em;
    margin-top: 0.8rem;
    margin-bottom: 1.8rem;
    text-transform: uppercase;
}
.callout {
    background: var(--light-bg);
    border-left: 4px solid var(--navy);
    padding: 0.9rem 1.1rem;
    margin: 0.5rem 0 1.2rem 0;
    border-radius: 0 8px 8px 0;
    font-size: 0.88rem;
    color: var(--t);
}
.section-header {
    background: var(--y);
    color: var(--navy);
    padding: 0.6rem 1rem;
    border-radius: 6px;
    font-weight: 800;
    font-size: 1rem;
    font-family: var(--heading);
    margin: 1.5rem 0 0.8rem 0;
    display: flex;
    align-items: center;
    gap: 0.4rem;
}
.small-meta {
    font-size: 0.78rem;
    color: var(--dim);
    margin-top: -0.2rem;
    margin-bottom: 0.5rem;
}
</style>
""",
    unsafe_allow_html=True,
)

# ─────────────────────────────────────
# STATE
# ─────────────────────────────────────
# FIX: unit_drafts 키를 zero-padded 형식("01"~"13")으로 통일
DEFAULT_STATE = {
    "style_dna": "",
    "merged_analysis": "",
    "gap_diagnosis": "",
    "story_reinforcement": {"기": "", "승": "", "전": "", "결": ""},
    "story_reinforcement_merged": "",
    "unit_blueprints": {
        "01-02": "",
        "03-04": "",
        "05-06": "",
        "07-08": "",
        "09-10": "",
        "11-12": "",
    },
    "unit_drafts": {f"{i:02d}" if i < 13 else "13": "" for i in range(1, 14)},
    "chapter_titles": {f"{i:02d}" if i < 13 else "13": "" for i in range(1, 14)},
    "ch1_stage_a": "",
    "ch1_stage_b": "",
    "ch1_stage_c": "",
    "unit_summaries": {},
    "quality_report": {},
    "character_tracker": {},
    "title_review": "",
    "status_message": "",
    "status_type": "info",
    "status_shown": False,   # v3.15 — 상태 메시지 1회 표시 후 소거용
    # v3.15 M17 Continuity Ledger — Unit별 물리 상태 원장 (문자열 키)
    "continuity_ledger": {},
    # v3.16 회차 제목 재검토 결과
    "chapter_title_review": "",
    # v3.1 시나리오 → 소설화 모드 상태
    "scenario_text": "",           # 업로드된 시나리오 원문
    "scenario_stats": {},          # 통계
    "scenario_extracted": {},      # Sonnet 추출 결과 (dict)
    "scenario_mapping_text": "",   # STEP 4 주입용 텍스트
    "scenario_fields_applied": False,  # STEP 1 자동 입력 완료 여부
    # v3.2 Creator Engine JSON → 소설화 모드 상태
    "creator_json_data": {},       # 업로드된 Creator JSON 원본 (dict)
    "creator_json_meta": {},       # UI 미리보기용 메타
    # v3.3 Idea Engine JSON → 소설화 모드 상태
    "idea_json_data": {},          # 업로드된 Idea JSON 원본 (dict)
    "idea_json_meta": {},          # UI 미리보기용 메타
    "idea_pending_items": [],      # 미결정 항목 목록
    # v3.4 STEP 1 입력 필드 — 세션 보존 + 프로젝트 저장 대상
    # (기존에는 위젯 지역변수로만 존재해 저장·복원이 불가능했다)
    "f_working_title": "",
    "f_genre": "",
    "f_format_mode": "장편소설",
    "f_pov": "3인칭 제한",
    "f_target_length": "",
    "f_style_strength": "중",
    "f_overview": "",
    "f_characters": "",
    "f_synopsis": "",
    "f_notes": "",
    "f_style_sample": "",
    "f_profession_protagonist": "",
    "f_profession_antagonist": "",
    "f_period_mode": "현대 (시대 주입 없음)",
    "f_period_labels": [],
    "f_locked_text": "",
    "f_open_text": "",
}

# 저장/불러오기 보조 키 (프로젝트 저장 대상 아님 — DEFAULT_STATE 밖)
if "_last_loaded_project" not in st.session_state:
    st.session_state["_last_loaded_project"] = ""
if "_last_truncated" not in st.session_state:
    st.session_state["_last_truncated"] = False

# v3.13 M15-B / v3.14 M16 — 집필 옵션 (STEP 4 설계보다 먼저 존재해야 한다)
if "metric_watchlist_on" not in st.session_state:
    st.session_state["metric_watchlist_on"] = True
if "signature_food_opening" not in st.session_state:
    st.session_state["signature_food_opening"] = True  # 작가 시그니처 — 기본 ON

for k, v in DEFAULT_STATE.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─────────────────────────────────────
# HELPERS
# ─────────────────────────────────────
def apply_extracted_to_fields(ex: dict) -> None:
    """추출 결과(scenario/creator/idea)를 STEP 1 위젯 세션 키에 직접 반영. (v3.4)

    위젯이 key 방식으로 바뀌면서 value= 인자를 쓸 수 없게 됐다.
    (key와 value를 함께 쓰면 재실행 시 session_state가 우선해 추출값이 무시된다)
    따라서 변환 직후 이 함수로 세션 값을 직접 채운 뒤 rerun 한다.
    """
    if not isinstance(ex, dict):
        return
    mapping = {
        "f_working_title": (ex.get("logline") or "")[:30],
        "f_genre": ex.get("genre") or "",
        "f_overview": ex.get("overview") or "",
        "f_characters": ex.get("characters") or "",
        "f_synopsis": ex.get("synopsis") or "",
        "f_notes": ex.get("notes") or "",
        "f_profession_protagonist": ex.get("profession_protagonist") or "",
        "f_profession_antagonist": ex.get("profession_antagonist") or "",
        "f_locked_text": ex.get("locked_text") or "",
        "f_open_text": ex.get("open_text") or "",
    }
    for k, v in mapping.items():
        if v:
            st.session_state[k] = v

    # 시대 키 → 수동 선택 모드 + 라벨 반영
    pk = ex.get("period_keys") or []
    if pk and PERIOD_KEYS:
        labels = []
        for key in pk[:2]:
            if key in PERIOD_KEYS:
                labels.append(f"{key} · {get_period_label(key)}")
        if labels:
            st.session_state["f_period_mode"] = "수동 선택"
            st.session_state["f_period_labels"] = labels


# 프로젝트 저장 대상 키 — DEFAULT_STATE 전체를 저장한다.
# 저장에서 제외할 일시적 키만 명시한다.
_SAVE_EXCLUDE_KEYS = {
    "status_message",
    "status_type",
    "status_shown",   # v3.15
}


def build_project_snapshot() -> dict:
    """현재 작업 상태 전체를 저장용 dict로 만든다. (v3.4)

    사고 패턴 B(통합본 동기화 사고) 방지 — 캐시가 아니라
    session_state의 최신 값을 직접 참조한다.
    """
    data = {}
    for k in DEFAULT_STATE.keys():
        if k in _SAVE_EXCLUDE_KEYS:
            continue
        data[k] = st.session_state.get(k, DEFAULT_STATE[k])

    # 저장 시점 메타
    _ex = st.session_state.get("scenario_extracted", {}) or {}
    return {
        "_novel_engine": {
            "engine": "Novel Engine",
            "version": NOVEL_ENGINE_VERSION,
            "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "title": st.session_state.get("f_working_title", "") or "(제목 미정)",
            "source": _ex.get("_source", "manual"),
            "source_title": _ex.get("_source_title", ""),
        },
        "state": data,
    }


def restore_project_snapshot(payload: dict) -> tuple:
    """저장된 프로젝트를 session_state에 복원한다. (v3.4)

    Returns:
        (성공여부, 메시지, 메타dict)
    """
    if not isinstance(payload, dict):
        return False, "파일 형식이 올바르지 않습니다.", {}

    meta = payload.get("_novel_engine", {})
    state = payload.get("state")
    if not isinstance(state, dict):
        return False, (
            "Novel Engine 프로젝트 파일이 아닙니다. "
            "Creator/Idea JSON이라면 STEP 0의 해당 탭을 사용하세요."
        ), {}

    restored, skipped = 0, 0
    for k, default in DEFAULT_STATE.items():
        if k in _SAVE_EXCLUDE_KEYS:
            continue
        if k not in state:
            skipped += 1
            continue
        v = state[k]
        # 타입이 크게 어긋나면 기본값 유지 (구버전 파일 호환)
        if default is not None and v is not None:
            if isinstance(default, dict) and not isinstance(v, dict):
                skipped += 1
                continue
            if isinstance(default, list) and not isinstance(v, list):
                skipped += 1
                continue
            if isinstance(default, bool) and not isinstance(v, bool):
                skipped += 1
                continue
        st.session_state[k] = v
        restored += 1

    msg = f"{restored}개 항목 복원"
    if skipped:
        msg += f" ({skipped}개는 구버전/형식 불일치로 기본값 유지)"
    return True, msg, meta


def count_written_units() -> int:
    """실제 원고가 있는 Unit 수."""
    drafts = st.session_state.get("unit_drafts", {}) or {}
    return sum(1 for v in drafts.values() if isinstance(v, str) and v.strip())


def get_client() -> Optional["anthropic.Anthropic"]:
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if not api_key or anthropic is None:
        return None
    return anthropic.Anthropic(api_key=api_key)


def looks_truncated(text: str) -> bool:
    """결과가 문장 중간에서 끊겼는지 판정한다. (v3.4.2)

    stop_reason을 받지 못하는 경로도 있으므로 텍스트 자체로도 확인한다.
    한국어 산문 기준으로 정상 종료는 보통 마침표·물음표·느낌표·따옴표·
    목록기호 등으로 끝난다.
    """
    if not text:
        return False
    tail = text.rstrip()
    if not tail:
        return False
    # 정상 종료로 볼 수 있는 마지막 글자
    ok_endings = ('.', '!', '?', '"', "'", '”', '’', ')', ']', '}',
                  '다', '음', '함', '임', '것', '요', '⟩', '>', '…', '—', ':', '·')
    if tail.endswith(ok_endings):
        return False
    # 목록/표/헤딩 줄로 끝나면 구조적 종료로 보고 정상 처리
    last_line = tail.split('\n')[-1].strip()
    if last_line.startswith(('#', '-', '*', '|', '>', '•')):
        return False
    # 숫자 목록 (1. 2. ①② 등)
    if re.match(r'^\s*(\d+[.)]|[①-⑳])', last_line):
        return False
    return True


def response_text(response) -> str:
    """Anthropic 응답에서 텍스트 블록만 골라 결합한다. (v3.3.2)

    adaptive thinking이 켜진 모델(Sonnet 5, Opus 4.8 등)은 첫 블록이
    ThinkingBlock일 수 있어 content[0].text 접근이 실패한다.
    """
    if response is None or not getattr(response, "content", None):
        return ""
    parts = []
    for block in response.content:
        if getattr(block, "type", None) == "text":
            t = getattr(block, "text", "")
            if t:
                parts.append(t)
    if parts:
        return "\n".join(parts).strip()
    for block in response.content:
        t = getattr(block, "text", None)
        if isinstance(t, str) and t.strip():
            return t.strip()
    return ""


def llm_call(user_prompt: str, max_tokens: int = MAX_TOKENS_MID, use_opus: bool = False) -> str:
    client = get_client()
    if client is None:
        return (
            "[오프라인 미리보기 모드]\n\n"
            "ANTHROPIC_API_KEY가 설정되지 않았거나 anthropic 패키지가 설치되지 않았습니다.\n"
            "실제 모델 호출 대신 프롬프트 초안만 구성된 상태입니다.\n\n"
            + user_prompt[:4000]
        )

    model = MODEL_OPUS if use_opus else DEFAULT_MODEL

    # v3.3.4 — 스트리밍 호출.
    # max_tokens가 크면 SDK가 non-streaming 호출을 거부한다
    # ('Streaming is required for operations that may take longer than 10 minutes').
    # 본문 집필도 장문이라 스트리밍으로 통일한다.
    result_text = ""
    stop_reason = ""
    try:
        parts = []
        with client.messages.stream(
            model=model,
            max_tokens=max_tokens,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}],
        ) as stream:
            for chunk in stream.text_stream:
                parts.append(chunk)
            final = stream.get_final_message()
            stop_reason = getattr(final, "stop_reason", "") or ""
            joined = "".join(parts).strip()
            result_text = joined if joined else response_text(final)
    except (AttributeError, TypeError):
        # 구버전 SDK 폴백
        response = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}],
        )
        result_text = response_text(response)
        stop_reason = getattr(response, "stop_reason", "") or ""

    # v3.4.1 — 빈 응답을 조용히 통과시키지 않는다.
    # 예전에는 ""를 그대로 반환해 "완료" 메시지만 뜨고 내용은 비는
    # 사고가 났다(설계·원고가 저장은 됐는데 화면에 안 보임).
    if not (result_text or "").strip():
        raise RuntimeError(
            "모델이 빈 응답을 반환했습니다. 잠시 후 다시 시도하거나, "
            "입력 자료가 너무 길지 않은지 확인하세요."
        )

    # v3.4.2 — 잘림 감지.
    # 토큰 상한에 걸려 문장 중간에서 끊긴 결과를 작가가 눈으로 발견할 때까지
    # 모르는 사고가 있었다. 이제 즉시 경고를 띄운다.
    if stop_reason == "max_tokens":
        st.session_state["_last_truncated"] = True
        st.warning(
            f"⚠️ 응답이 최대 길이({max_tokens:,} 토큰)에 도달해 **끝부분이 잘렸습니다.** "
            "결과 마지막이 문장 중간에서 끝났는지 확인하세요. "
            "입력 자료(작품 개요·캐릭터·줄거리)를 줄이거나 다시 시도하면 개선될 수 있습니다."
        )
    else:
        st.session_state["_last_truncated"] = False

    return result_text


def merge_nonempty(parts: List[str], sep: str = "\n\n") -> str:
    return sep.join([p.strip() for p in parts if p and p.strip()])


def ensure_final_ending(text: str, unit_no: int) -> str:
    text = (text or "").rstrip()
    if unit_no in (12, 13) and not text.endswith("끝."):
        return f"{text}\n\n끝."
    return text


def is_incomplete_text(text: str, unit_no: int) -> bool:
    txt = (text or "").strip()
    if not txt:
        return True

    min_len = UNIT_MIN_LENGTHS.get(unit_no, 6000)
    if len(txt) < min_len:
        return True

    valid_endings = [".", "!", "?", "\u201d", "\"", "'", "\u2019", "끝."]
    if not any(txt.endswith(e) for e in valid_endings):
        return True

    lines = [line for line in txt.splitlines() if line.strip()]
    if len(lines) < 12:
        return True

    return False


# ─────────────────────────────────────
# 품질 자동 체크 (생성 후 즉시 경고) — v3.0 M1 BJND Scene Enforcer 연동
# ─────────────────────────────────────
def analyze_unit_quality(text: str) -> dict:
    """생성된 Unit 원고의 품질 문제를 자동 감지한다.
    v3.0: BJND 임계치 강화 + 자동 재생성용 violations 구조화 반환.

    Returns:
        {
            "issues": [경고 문자열 리스트 — UI 표시용],
            "stats": {지표명: 값},
            "violations": {지표키: {count, threshold, severity}},  # v3.0 신규
            "should_regenerate": bool,  # v3.0 신규 — 자동 재생성 트리거 여부
        }
    """
    import re
    if not text or not text.strip():
        return {"issues": [], "stats": {}, "violations": {}, "should_regenerate": False}

    issues = []
    stats = {}
    violations = {}  # v3.0 신규: 재생성 트리거용

    # ── 종결어미 반복 (v3.0: "있었다" 15→10, "것이었다" 3→2) ──
    cnt_isseotda = len(re.findall(r"있었다", text))
    cnt_geosieotda = len(re.findall(r"것이었다", text))
    stats["있었다"] = cnt_isseotda
    stats["것이었다"] = cnt_geosieotda

    if cnt_isseotda > BJND_THRESHOLDS["있었다"]:
        issues.append(f"⚠️ '있었다' {cnt_isseotda}회 — {BJND_THRESHOLDS['있었다']}회 이하로 줄이세요. 구체 동사로 대체.")
        violations["있었다"] = {
            "count": cnt_isseotda,
            "threshold": BJND_THRESHOLDS["있었다"],
            "severity": "high",
        }
    if cnt_geosieotda > BJND_THRESHOLDS["것이었다"]:
        issues.append(f"⚠️ '것이었다' 해설체 {cnt_geosieotda}회 — {BJND_THRESHOLDS['것이었다']}회 이하로.")
        violations["것이었다"] = {
            "count": cnt_geosieotda,
            "threshold": BJND_THRESHOLDS["것이었다"],
            "severity": "high",
        }

    # ── 대사 태그 반복 ──
    cnt_said = len(re.findall(r"말했다", text))
    cnt_asked = len(re.findall(r"물었다", text))
    cnt_answered = len(re.findall(r"대답했다", text))
    tag_total = cnt_said + cnt_asked + cnt_answered
    stats["대사태그 합계"] = tag_total
    if tag_total > BJND_THRESHOLDS["대사태그"]:
        issues.append(f"⚠️ 대사 태그(말했다/물었다/대답했다) {tag_total}회 — 행동 태그로 대체하세요.")
        violations["대사태그"] = {
            "count": tag_total,
            "threshold": BJND_THRESHOLDS["대사태그"],
            "severity": "medium",
        }

    # ── "마치 ~처럼" / "~듯했다" / "~같았다" 반복 (v3.0 신규) ──
    cnt_macheoreom = len(re.findall(r"마치\s*.*?처럼", text))
    cnt_deuthaetda = len(re.findall(r"듯했다|듯이", text))
    cnt_gatatda = len(re.findall(r"같았다", text))
    simile_total = cnt_macheoreom + cnt_deuthaetda + cnt_gatatda
    stats["비유 패턴"] = simile_total
    # 문단 단위는 정확 측정이 어려우니, Unit 전체에서 Unit당 약 1문단 1회 기준으로 임계
    # 4개 문단당 1개 정도가 리미트 → 대략 Unit 분량 기준 6~8회까지 허용
    if simile_total > 8:
        issues.append(f"⚠️ '마치 ~처럼/듯했다/같았다' 합계 {simile_total}회 — 비유 과잉.")
        violations["마치처럼"] = {
            "count": simile_total,
            "threshold": 8,
            "severity": "medium",
        }

    # ── 장면 반복 ──
    cnt_phone = len(re.findall(r"전화|휴대폰이|진동했다|문자|메시지가", text))
    cnt_window = len(re.findall(r"창밖|창 밖|유리창|내려다보", text))
    cnt_elevator = len(re.findall(r"엘리베이터|로비를|현관을", text))
    stats["전화/메시지"] = cnt_phone
    stats["창밖 묘사"] = cnt_window
    if cnt_phone > 4:
        issues.append(f"⚠️ 전화/메시지 장면 {cnt_phone}회 — 대면/발견/관찰로 대체하세요.")
        violations["전화"] = {"count": cnt_phone, "threshold": 4, "severity": "medium"}
    if cnt_window > 3:
        issues.append(f"⚠️ '창밖' 묘사 {cnt_window}회 — 다른 방식으로 인물 내면을 쓰세요.")
        violations["창밖"] = {"count": cnt_window, "threshold": 3, "severity": "low"}
    if cnt_elevator > 2:
        issues.append(f"⚠️ 엘리베이터/로비 {cnt_elevator}회 — 이동 묘사를 줄이세요.")
        violations["엘리베이터"] = {"count": cnt_elevator, "threshold": 2, "severity": "low"}

    # ── 시제 체크 (v3.0: 치명적) ──
    present_patterns = re.findall(
        r"(?:한다|된다|이다|간다|온다|본다|듣는다|만든다|열린다|닫힌다|울린다|채운다|넣는다|흐른다)\.",
        text,
    )
    cnt_present = len(present_patterns)
    stats["현재형 종결"] = cnt_present
    if cnt_present > BJND_THRESHOLDS["현재형"]:
        issues.append(f"🚨 현재형 종결어미 {cnt_present}회 감지 — 소설은 과거형(~했다)으로 써야 합니다!")
        violations["현재형"] = {
            "count": cnt_present,
            "threshold": BJND_THRESHOLDS["현재형"],
            "severity": "critical",
        }

    # ── 계량 수치 (v3.12 M15 Metric Precision Ban) ──
    # ① 숫자 + 계량 단위  ② 소수점 숫자
    # 사회적 단위(년·월·일·시·분·원·억·만·명·개·번·층·살·회·화)는 제외한다.
    metric_unit_hits = re.findall(METRIC_UNIT_RE, text)
    decimal_hits = re.findall(METRIC_DECIMAL_RE, text)
    # 중복 계산 방지 — 소수점이 계량 단위와 함께 잡힌 경우는 한 번만 센다
    metric_total = len(set(metric_unit_hits)) + max(
        0, len(decimal_hits) - sum(1 for h in metric_unit_hits if "." in h)
    )
    stats["계량 수치"] = metric_total
    if metric_total > BJND_THRESHOLDS["계량수치"]:
        sample = ", ".join(list(dict.fromkeys(metric_unit_hits + decimal_hits))[:4])
        issues.append(
            f"🚨 계량 수치 {metric_total}회 감지 ({sample}) — "
            f"서술문의 계량 단위·소수점은 금지입니다. "
            f"몸의 단위(손끝·한 뼘·한 박자)나 비교로 치환하세요."
        )
        violations["계량수치"] = {
            "count": metric_total,
            "threshold": BJND_THRESHOLDS["계량수치"],
            "severity": "high",
        }

    # ── 접속부사 과잉 ──
    cnt_conj = len(re.findall(r"(?:그러나|하지만|그리고|또한|그래서|따라서)", text))
    stats["접속부사"] = cnt_conj
    if cnt_conj > 15:
        issues.append(f"⚠️ 접속부사 {cnt_conj}회 — 접속부사 없이 문장을 병치하세요.")
        violations["접속부사"] = {"count": cnt_conj, "threshold": 15, "severity": "low"}

    # ── 같은 행동 반복 ──
    action_patterns = [
        (r"포크를 내려놓", "포크를 내려놓"),
        (r"잔을 내려놓", "잔을 내려놓"),
        (r"눈을 감", "눈을 감"),
        (r"한숨을 쉬", "한숨을 쉬"),
        (r"고개를 끄덕", "고개를 끄덕"),
        (r"고개를 저", "고개를 저"),
    ]
    for pattern, label in action_patterns:
        cnt = len(re.findall(pattern, text))
        if cnt >= 3:
            issues.append(f"⚠️ '{label}' {cnt}회 반복 — 같은 동작을 줄이세요.")

    # ── 분량 ──
    stats["총 글자수"] = len(text)

    # ══════════════════════════════════════════════════════
    # v3.15 신규 진단 지표 3종
    # ══════════════════════════════════════════════════════

    # ① 신체 반응 어휘 반복 (M11 보강)
    #    실측 배경: 《사랑한다고 했잖아》 UNIT 01에서 "손끝" 7회,
    #    UNIT 02에서 3회. "손끝이 차가웠다/굳었다/얼어붙었다"가 한 Unit 안에서 돌았다.
    #    "심장이 늑골을 두드렸다"(01) ↔ "심장이 갈비뼈를 쳤다"(02)처럼
    #    같은 비유를 부위 이름만 바꿔 재사용하는 사고도 여기서 잡는다.
    body_hits = BODY_REACTION_RE.findall(text)
    cnt_body = len(body_hits)
    stats["신체반응 어휘"] = cnt_body
    if cnt_body > BJND_THRESHOLDS["신체반응"]:
        # 어떤 어휘가 몰렸는지 상위 3개를 함께 보여준다.
        top = sorted(
            {w: body_hits.count(w) for w in set(body_hits)}.items(),
            key=lambda kv: -kv[1],
        )[:3]
        detail = ", ".join(f"{w} {c}회" for w, c in top)
        issues.append(
            f"⚠️ 신체 반응 어휘 {cnt_body}회 ({detail}) — "
            f"{BJND_THRESHOLDS['신체반응']}회 이하로. 감정을 같은 부위로만 번역하지 마세요."
        )
        violations["신체반응"] = {
            "count": cnt_body,
            "threshold": BJND_THRESHOLDS["신체반응"],
            "severity": "medium",
        }

    # ② 감정 라벨링 / 기억 오류 장치 (M18)
    #    "위화감이라고 이름 붙이지는 않았다", "그 차이가 어디서 오는지 따져보지 않았다"
    #    — 그루밍 서사에서 강력한 장치지만 반복되면 서술자가 독자에게
    #    정답을 짚어주는 개입이 된다.
    cnt_label = len(NARRATOR_LABEL_RE.findall(text))
    stats["서술자 라벨링"] = cnt_label
    if cnt_label > BJND_THRESHOLDS["기억오류"]:
        issues.append(
            f"⚠️ 서술자 라벨링 구문 {cnt_label}회 — "
            f"{BJND_THRESHOLDS['기억오류']}회까지. 반복하면 기법이 노출됩니다. (M18)"
        )
        violations["기억오류"] = {
            "count": cnt_label,
            "threshold": BJND_THRESHOLDS["기억오류"],
            "severity": "medium",
        }

    # ③ 회상 진입 마커 (M12 / Chapter 1 Stage B 회상 금지)
    #    Chapter 1은 현재 시간대만 다루는데 "지난달에도 그랬다"류 회상 블록이
    #    Stage B에 들어간 사례가 있었다. 카운터가 없어 감지되지 않았다.
    cnt_fb = len(FLASHBACK_MARKER_RE.findall(text))
    stats["회상 마커"] = cnt_fb
    if cnt_fb > BJND_THRESHOLDS["회상마커"]:
        issues.append(
            f"⚠️ 회상 진입 마커 {cnt_fb}회 — 회상이 Unit 분량의 30%를 넘지 않는지 "
            f"확인하세요. Chapter 1이면 회상 자체를 재검토하세요. (M12)"
        )
        violations["회상마커"] = {
            "count": cnt_fb,
            "threshold": BJND_THRESHOLDS["회상마커"],
            "severity": "low",
        }

    # ④ 잔여 마커 검출 (v3.15) — 본문 중간 '끝.' / 마크다운 잔재
    #    Stage A·B가 '끝.'을 출력해 병합본 한가운데에 박힌 사고,
    #    첫 줄이 **"대사"** 형태로 나간 사고를 표면화한다.
    stray_end = len(re.findall(r"(?<!\A)^\s*끝\.\s*$", text.strip(), re.M))
    # 마지막 줄의 '끝.'은 정상(Unit 12/13)이므로 제외
    if text.strip().endswith("끝.") and stray_end > 0:
        stray_end -= 1
    cnt_md = len(re.findall(r"\*\*|^#{1,6}\s|^---\s*$", text, re.M))
    stats["잔여 마커"] = stray_end + cnt_md
    if stray_end > 0 or cnt_md > 0:
        parts = []
        if stray_end:
            parts.append(f"본문 중간 '끝.' {stray_end}개")
        if cnt_md:
            parts.append(f"마크다운 기호 {cnt_md}개")
        issues.append(
            "⚠️ 잔여 마커 검출 — " + " / ".join(parts) +
            ". [본문 정리] 버튼으로 제거할 수 있습니다."
        )
        violations["잔여마커"] = {
            "count": stray_end + cnt_md,
            "threshold": 0,
            "severity": "low",
        }

    # v3.0 신규: 자동 재생성 트리거 판정
    # severity가 "critical" 또는 "high"인 violation이 있으면 재생성 대상
    should_regenerate = any(
        v.get("severity") in ("critical", "high")
        for v in violations.values()
    )

    return {
        "issues": issues,
        "stats": stats,
        "violations": violations,
        "should_regenerate": should_regenerate,
    }


# =====================================================================
# v3.15 — 본문 정리기 (후처리)
# =====================================================================
def sanitize_manuscript(text: str, is_final_unit: bool = False) -> tuple:
    """생성 결과에서 마커 잔재를 걷어낸다. (v3.15)

    배경 — 두 가지 사고가 실제로 발생했다.
      ① Stage A·B가 각각 마지막 줄에 '끝.'을 출력해, Chapter 1 병합본
         한가운데에 '끝.'이 두 번 박혔다. (저장 JSON에서 직접 확인)
      ② UNIT 02 첫 줄이 **"얼굴이 왜 그래."** 형태로 나가 DOCX에
         별표가 그대로 들어갔다.

    ★ 작업 규칙 7 준수 — 본문 내용은 건드리지 않는다. ★
    제거 대상은 '원고가 아닌 마커'뿐이다. 문장, 어휘, 문단 구조는
    한 글자도 바꾸지 않는다.

    Args:
        text: 원본 텍스트
        is_final_unit: True면 맨 마지막 줄의 '끝.'은 보존한다 (Unit 12/13)

    Returns:
        (정리된 텍스트, 변경 내역 리스트)
    """
    if not text or not text.strip():
        return text, []

    log = []
    out = text

    # ── ① 강조 마크다운 제거 (**bold** / *italic* 감싸기) ──
    #     별표 안의 내용은 그대로 두고 별표만 벗긴다.
    n_bold = len(re.findall(r"\*\*(.+?)\*\*", out, re.S))
    if n_bold:
        out = re.sub(r"\*\*(.+?)\*\*", r"\1", out, flags=re.S)
        log.append(f"굵게 표시 마크다운 {n_bold}개 제거")

    # ── ①-2 기울임 마크다운 제거 (v3.16.3) ──
    #     실제 사고 — 문자 메시지가 *오늘도 늦어? 밥은 먹었고.* 형태로 나갔다.
    #     장면 전환 마커(별표만 있는 줄)와 구분하기 위해, 별표 사이에
    #     내용이 있는 경우만 벗긴다.
    n_ital = len(re.findall(r"(?<!\*)\*(?!\s)([^*\n]{1,200}?)(?<!\s)\*(?!\*)", out))
    if n_ital:
        out = re.sub(r"(?<!\*)\*(?!\s)([^*\n]{1,200}?)(?<!\s)\*(?!\*)", r"\1", out)
        log.append(f"기울임 마크다운 {n_ital}개 제거")

    # ── ② 헤딩·수평선 마커 제거 (줄 전체가 마커인 경우만) ──
    n_hr = len(re.findall(r"^\s*(?:---+|===+|\*\*\*+)\s*$", out, re.M))
    if n_hr:
        out = re.sub(r"^\s*(?:---+|===+|\*\*\*+)\s*$", "", out, flags=re.M)
        log.append(f"수평선 마커 {n_hr}개 제거")

    # 헤딩은 [CHAPTER n] 첫 줄을 건드리면 안 되므로 '#' 접두만 벗긴다.
    n_head = len(re.findall(r"^\s*#{1,6}\s+", out, re.M))
    if n_head:
        out = re.sub(r"^\s*#{1,6}\s+", "", out, flags=re.M)
        log.append(f"헤딩 기호 {n_head}개 제거")

    # ── ③ 본문 중간의 '끝.' 제거 ──
    lines = out.split("\n")
    # 마지막 비어있지 않은 줄의 인덱스
    last_idx = -1
    for i in range(len(lines) - 1, -1, -1):
        if lines[i].strip():
            last_idx = i
            break

    removed_end = 0
    kept_final = False
    for i, ln in enumerate(lines):
        if ln.strip() in ("끝.", "끝", "— 끝 —", "-끝-"):
            if is_final_unit and i == last_idx:
                kept_final = True
                continue
            lines[i] = ""
            removed_end += 1
    if removed_end:
        out = "\n".join(lines)
        note = f"본문 중간 '끝.' {removed_end}개 제거"
        if kept_final:
            note += " (마지막 줄 '끝.'은 보존)"
        log.append(note)

    # ── ④ 집필 메타 발화 검출 (v3.16.2) ──
    #     "Stage B가 닫혔으니 ~ Stage C를 열겠습니다" 같은 문장이 본문에
    #     섞여 나온 사고가 있었다. 모델이 작가에게 하는 말을 원고에 쓴 것이다.
    #     ★ 본문 문장이므로 자동 삭제하지 않는다 (작업 규칙 7). 로그로만 알린다. ★
    meta_hits = re.findall(
        r"^.*(?:Stage\s*[ABC]|스테이지\s*[ABCabc])(?=.*(?:겠습니다|하겠|보여드|이어서)).*$",
        out, re.M,
    )
    meta_hits += re.findall(
        r"^.*(?:열겠습니다|그리겠습니다|작성하겠습니다|보여드리겠습니다|이어서 쓰겠습니다).*$",
        out, re.M,
    )
    meta_hits = [h.strip() for h in dict.fromkeys(meta_hits) if h.strip()]
    if meta_hits:
        log.append(
            f"⚠️ 집필 메타 발화로 보이는 문장 {len(meta_hits)}건 발견 "
            f"(자동 삭제하지 않음) — \"{meta_hits[0][:40]}…\""
        )

    # ── ⑤ 3줄 이상 연속 공백을 2줄로 정리 ──
    #     위 제거 과정에서 생긴 빈 줄만 정돈한다.
    if re.search(r"\n{4,}", out):
        out = re.sub(r"\n{4,}", "\n\n\n", out)
        log.append("과다 공백 줄 정리")

    return out.strip(), log


def build_retry_hint(violations: dict) -> str:
    """v3.0 M1: 자동 재생성 시 프롬프트에 주입할 위반 지표 힌트를 생성.

    prompt.py의 build_unit_draft_prompt에 retry_hint 인자로 전달된다.
    """
    if not violations:
        return ""

    # severity 순서대로 정렬 (critical → high → medium → low)
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    sorted_violations = sorted(
        violations.items(),
        key=lambda kv: severity_order.get(kv[1].get("severity", "low"), 9)
    )

    lines = []
    for key, info in sorted_violations:
        count = info.get("count", 0)
        threshold = info.get("threshold", 0)
        severity = info.get("severity", "")

        # 지표별 구체 지시문 생성
        if key == "있었다":
            target = max(threshold - 2, 5)  # 여유 -2
            lines.append(
                f"- '있었다' 사용: 이전 {count}회 (임계 {threshold}회 초과). "
                f"이번엔 {target}회 이하로 작성. 구체 동사로 전환: '놓여 있었다'→'놓였다', '서 있었다'→'기대어 있었다' 등."
            )
        elif key == "것이었다":
            lines.append(
                f"- '것이었다' 해설체: 이전 {count}회 (임계 {threshold}회 초과). "
                f"이번엔 1회 이하로. 서술 구조를 재설계하라."
            )
        elif key == "대사태그":
            target = max(threshold - 2, 8)
            lines.append(
                f"- 대사 태그 합계(말했다/물었다/대답했다): 이전 {count}회 (임계 {threshold}회 초과). "
                f"이번엔 {target}회 이하로. 대사 10개 중 6개 이상은 행동 태그로 대체."
            )
        elif key == "마치처럼":
            lines.append(
                f"- 비유 패턴(마치~처럼/듯했다/같았다): 이전 {count}회 (임계 {threshold}회 초과). "
                f"이번엔 5회 이하로. 문단당 1회만 허용."
            )
        elif key == "현재형":
            lines.append(
                f"- 🚨 현재형 종결: 이전 {count}회 (임계 {threshold}회 초과, 치명적). "
                f"이번엔 0회로. 모든 서술은 과거형(~했다, ~였다, ~었다)."
            )
        elif key == "전화":
            lines.append(
                f"- 전화/메시지 장면: 이전 {count}회 (임계 {threshold}회 초과). "
                f"이번엔 2회 이하로. 대면·발견·관찰로 대체."
            )
        elif key == "창밖":
            lines.append(
                f"- '창밖' 묘사: 이전 {count}회 (임계 {threshold}회 초과). "
                f"이번엔 1회 이하로. 걷기·운전·사물 다루기 등 다른 방식으로 내면 표현."
            )
        elif key == "엘리베이터":
            lines.append(
                f"- 엘리베이터/로비/현관 이동: 이전 {count}회 (임계 {threshold}회 초과). "
                f"이번엔 1회 이하로. 장면 전환으로 처리."
            )
        elif key == "계량수치":
            lines.append(
                f"- 🚨 계량 수치: 이전 {count}회 (임계 {threshold}회 초과). "
                f"이번엔 0회로. 서술문에서 계량 단위(센티미터·밀리미터·그램·도·퍼센트·초)와 "
                f"소수점 숫자를 전부 삭제하라. "
                f"'0.5센티미터씩 정렬했다'→'글자가 한 줄로 읽힐 때까지 손끝으로 밀었다', "
                f"'3밀리미터 어긋나면'→'비뚜름하면'. "
                f"인물의 정밀함은 숫자가 아니라 남이 못 보는 어긋남을 보는 눈으로 써라."
            )
        elif key == "접속부사":
            lines.append(
                f"- 접속부사 과잉: 이전 {count}회. 이번엔 10회 이하로. 문장 병치로 관계를 암시."
            )
        else:
            lines.append(f"- {key}: 이전 {count}회 (임계 {threshold}회 초과). 감소 필요.")

    return "\n".join(lines)



# ─────────────────────────────────────
# Unit 요약 자동 생성
# ─────────────────────────────────────
def analyze_cross_unit_repetition(unit_no: int, text: str, window: int = 3) -> list:
    """직전 Unit들과 겹치는 표현을 찾아낸다. (v3.15.1)

    배경 — analyze_unit_quality()는 한 Unit 안의 반복만 본다.
    실제 사고는 Unit을 건너뛰며 났다.
      · "심장이 늑골을 두드렸다"(UNIT 01) ↔ "심장이 갈비뼈를 쳤다"(UNIT 02)
        — 늑골과 갈비뼈는 같은 말이다. 부위 이름만 바꾼 동일 비유 재사용.
      · "손끝이 차가웠다 / 굳었다 / 얼어붙었다"가 두 Unit에 걸쳐 반복.
    독자는 Unit 경계를 의식하지 않으므로, 이어 읽으면 즉시 눈에 걸린다.

    ★ 진단만 한다. 본문은 고치지 않는다 (작업 규칙 7). ★

    Args:
        unit_no: 검사 대상 Unit 번호
        text: 검사 대상 Unit 본문
        window: 몇 개 앞 Unit까지 대조할지

    Returns:
        경고 문자열 리스트
    """
    if not text or not text.strip():
        return []

    drafts = st.session_state.get("unit_drafts", {}) or {}
    prev_texts = []
    for n in range(int(unit_no) - 1, max(0, int(unit_no) - 1 - window), -1):
        t = drafts.get(f"{n:02d}") or drafts.get(str(n)) or drafts.get(n) or ""
        if isinstance(t, str) and t.strip():
            prev_texts.append((n, t))
    if not prev_texts:
        return []

    warnings = []

    # ── ① 동일 비유 재사용 — 같은 신체 부위군을 같은 방식으로 다시 쓰는 경우 ──
    # 같은 뜻의 다른 표기를 하나로 묶는다. 늑골 = 갈비뼈, 뒷목 = 목덜미.
    synonym_groups = {
        "갈비뼈/늑골": r"(?:늑골|갈비뼈)",
        "목덜미/뒷목": r"(?:목덜미|뒷목)",
        "손끝": r"(?:손끝|손가락 끝)",
        "등골": r"(?:등골|등이 서늘|등이 차가)",
        "심장": r"심장",
        "명치": r"명치",
    }
    for label, pat in synonym_groups.items():
        cur = len(re.findall(pat, text))
        if cur == 0:
            continue
        for n, pt in prev_texts:
            prev_cnt = len(re.findall(pat, pt))
            if prev_cnt >= 2 and cur >= 2:
                warnings.append(
                    f"🔁 '{label}' — UNIT {n:02d}에서 {prev_cnt}회, "
                    f"이번 Unit에서 {cur}회. 같은 부위로 감정을 반복 번역하고 있습니다."
                )
                break

    # ── ② 특징적 어구 재사용 — 5~12자 어구가 앞 Unit에 그대로 있는 경우 ──
    # 흔한 관용구를 걸러내기 위해 조사·서술어만으로 된 조각은 제외한다.
    stop = ("그리고", "그러나", "그런데", "하지만", "그래서", "지윤은", "도균은")
    phrases = set()
    for m in re.finditer(r"[가-힣]{2,}(?:\s[가-힣]{2,}){1,2}", text):
        ph = m.group(0)
        if 6 <= len(ph) <= 14 and not ph.startswith(stop):
            phrases.add(ph)

    dup = []
    for n, pt in prev_texts:
        for ph in phrases:
            if ph in pt and len(dup) < 6:
                dup.append((ph, n))
    if dup:
        shown = ", ".join(f"'{p}'(UNIT {n:02d})" for p, n in dup[:4])
        warnings.append(
            f"🔁 앞 Unit과 겹치는 어구 {len(dup)}건 — {shown}"
            + (" 외" if len(dup) > 4 else "")
            + ". 의도한 모티프가 아니면 표현을 바꿔주세요."
        )

    return warnings


def generate_continuity_ledger(unit_no: int, text: str) -> str:
    """완성된 Unit에서 연속성 상태 원장을 추출해 저장한다. (v3.15 M17)

    unit_summaries는 '줄거리 요약'이라 물리 상태를 추적하지 못한다.
    실제 사고 — UNIT 01에서 지윤이 앞치마를 '입고' 있었는데 UNIT 02가
    "앞치마는 접어 가방에 넣은 채였다"로 시작했고, 그 한 줄이 UNIT 02
    서스펜스 전체의 근거라 챕터 논리가 무너졌다.

    원장은 시각·날씨·위치·착의·소지품·신체·통신·인지·독자·미회수
    10개 항목으로, 다음 Unit 집필 프롬프트에 HARD 블록으로 주입된다.
    """
    client = get_client()
    if not client or not text or not text.strip():
        return ""
    try:
        prev = gather_continuity_ledger(before_unit=unit_no)
        prompt = build_continuity_ledger_prompt(unit_no, text[:12000], prev)
        # v3.16.3 — max_tokens 1500은 부족했다. 실제 사고: UNIT 01 원장이
        # [착의] 항목 중간에서 잘려, [소지품]·[인지]·[통신] 등 뒤쪽 6개 항목이
        # 통째로 사라졌다. 그래서 "명함을 택시 안에서 받았다"는 사실이
        # 원장에 없었고, UNIT 02가 "파티 케이터링 화장실 앞에서 받았다"로
        # 왜곡해도 걸러지지 않았다.
        resp = client.messages.create(
            model=DEFAULT_MODEL,
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}],
        )
        ledger = response_text(resp).strip()
        if not ledger:
            st.warning(
                f"UNIT {unit_no:02d} 상태 원장 생성이 비어 있습니다. "
                "다음 Unit에서 연속성 검증이 작동하지 않습니다."
            )
            return ""

        # v3.16.3 — 잘림 검사. 원장이 잘리면 뒤쪽 항목이 통째로 사라진다.
        required = ["[시각·경과]", "[핵심 사실]", "[소지품]", "[인지]", "[미회수]"]
        missing = [k for k in required if k not in ledger]
        if missing or looks_truncated(ledger):
            st.warning(
                f"UNIT {unit_no:02d} 상태 원장이 불완전합니다 — "
                + (f"누락 항목 {', '.join(missing)}. " if missing else "문장 중간에서 끊겼습니다. ")
                + "STEP 5의 [점검 C · 상태 원장 생성]으로 다시 만들거나, "
                "원장 편집창에서 직접 채워주세요."
            )

        if "continuity_ledger" not in st.session_state:
            st.session_state["continuity_ledger"] = {}
        st.session_state["continuity_ledger"][f"{unit_no:02d}"] = ledger
        return ledger
    except Exception as e:
        st.warning(
            f"UNIT {unit_no:02d} 상태 원장 생성에 실패했습니다 ({type(e).__name__}). "
            "연속성 검증 없이 다음 Unit이 생성됩니다."
        )
        return ""


def gather_continuity_ledger(before_unit: int) -> str:
    """before_unit 직전까지의 상태 원장을 모아 반환한다. (v3.15 M17)

    사고 패턴 A 대응 — session_state는 정수 키, JSON 로드는 문자열 키라서
    양쪽 표기를 모두 조회한다.

    가장 최근 원장 1개를 우선으로 쓰되, 그 앞 원장 1개까지 함께 넘겨
    2단계 전의 미회수 복선도 놓치지 않게 한다.
    """
    led = st.session_state.get("continuity_ledger", {}) or {}
    if not led:
        return ""
    picked = []
    for n in range(int(before_unit) - 1, 0, -1):
        txt = led.get(f"{n:02d}") or led.get(str(n)) or led.get(n) or ""
        if isinstance(txt, str) and txt.strip():
            picked.append((n, txt.strip()))
        if len(picked) >= 2:
            break
    if not picked:
        return ""
    picked.reverse()
    return "\n\n".join(f"[UNIT {n:02d} 종료 시점 원장]\n{t}" for n, t in picked)


def resolve_chapter_title(unit_no: int, model_title: str, blueprints_text: str) -> tuple:
    """회차 제목을 확정한다. 설계안 제목이 정본이다. (v3.16)

    배경 — v3.15까지 제목이 두 군데서 만들어졌다.
      · STEP 4 설계의 '- 제목:' 필드 → "다정함의 온도"
      · STEP 5 집필의 '[CHAPTER n] — 서브타이틀' 지시 → 별개의 제목
    둘이 연결돼 있지 않아 어긋났고, 실제로는 모델이 집필 단계 지시를
    무시해서 UNIT 01·02 원고에 [CHAPTER n] 줄이 아예 없었다.

    v3.16부터 집필 프롬프트는 번호만 출력하고, 제목은 여기서 붙인다.
    설계안 제목은 2 Unit씩 전체 맥락에서 뽑혀 배열까지 돼 있으므로
    (04 "정확히 그 자리" → 05 "너무 정확한 밤") 그대로 정본으로 쓴다.

    Args:
        unit_no: Unit 번호
        model_title: 모델이 출력한 첫 줄 (없으면 "")
        blueprints_text: 전체 설계 텍스트

    Returns:
        (확정 제목, 안내 메시지) — 안내가 필요 없으면 메시지는 ""
    """
    n = int(unit_no)
    bp_title = extract_blueprint_chapter_title(blueprints_text, n)

    # 모델이 붙인 서브타이틀이 있으면 떼어낸다 (v3.16 이전 원고 호환)
    model_sub = ""
    if model_title:
        m = re.match(r"^\[CHAPTER[^\]]*\]\s*[—\-–]\s*(.+)$", model_title.strip())
        if m:
            model_sub = m.group(1).strip()

    if bp_title:
        title = f"[CHAPTER {n}] — {bp_title}"
        note = ""
        # 모델이 다른 제목을 지어냈다면 작가에게 알린다. 설계안이 이긴다.
        if model_sub and model_sub != bp_title:
            note = (
                f"집필 단계에서 '{model_sub}'라는 제목이 나왔지만, "
                f"설계안 제목을 정본으로 적용했습니다 — '{bp_title}'."
            )
        return title, note

    # 설계안에 제목이 없는 경우 — 모델 제목을 살리고, 그것도 없으면 번호만
    if model_sub:
        return f"[CHAPTER {n}] — {model_sub}", ""
    return f"[CHAPTER {n}]", (
        f"UNIT {n:02d} 설계안에 제목이 없어 번호만 넣었습니다. "
        "STEP 7에서 회차 제목을 채울 수 있습니다."
    )


def generate_unit_summary(unit_no: int, text: str) -> str:
    """완성된 Unit의 1줄 요약을 생성한다."""
    client = get_client()
    if not client or not text or not text.strip():
        return ""
    try:
        resp = client.messages.create(
            model=os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-5"),
            max_tokens=200,
            messages=[{
                "role": "user",
                "content": (
                    f"다음은 소설의 UNIT {unit_no:02d} 원고이다. "
                    "이 Unit의 핵심 사건, 인물 변화, 감정 상태를 1~2문장으로 요약하라. "
                    "요약만 출력하고 다른 말은 하지 마라.\n\n"
                    f"{text[:3000]}"
                ),
            }],
        )
        return response_text(resp).strip()
    except Exception:
        return ""


def gather_all_summaries() -> str:
    """모든 Unit 요약을 모아 반환한다."""
    summaries = st.session_state.get("unit_summaries", {})
    lines = []
    for i in range(1, 14):
        key = f"{i:02d}" if i < 13 else "13"
        s = summaries.get(key, "")
        if s:
            lines.append(f"[UNIT {key} 요약] {s}")
    return "\n".join(lines)


# ─────────────────────────────────────
# 캐릭터 등장 추적
# ─────────────────────────────────────
def extract_characters_from_text(text: str) -> str:
    """Unit 원고에서 등장 인물 목록을 추출한다."""
    client = get_client()
    if not client or not text or not text.strip():
        return ""
    try:
        resp = client.messages.create(
            model=os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-5"),
            max_tokens=300,
            messages=[{
                "role": "user",
                "content": (
                    "다음 소설 원고에서 등장하는 인물의 이름만 쉼표로 구분해서 나열하라. "
                    "이름만 출력하고 다른 말은 하지 마라. 이름이 없으면 '없음'.\n\n"
                    f"{text[:4000]}"
                ),
            }],
        )
        return response_text(resp).strip()
    except Exception:
        return ""


def track_characters(unit_key: str, text: str):
    """캐릭터 등장을 추적하고 세션에 저장한다."""
    if "character_tracker" not in st.session_state:
        st.session_state["character_tracker"] = {}
    names_str = extract_characters_from_text(text)
    if names_str and names_str != "없음":
        names = [n.strip() for n in names_str.split(",") if n.strip()]
        st.session_state["character_tracker"][unit_key] = names


def get_character_report() -> dict:
    """캐릭터 등장 추적 리포트를 생성한다."""
    tracker = st.session_state.get("character_tracker", {})
    if not tracker:
        return {}
    first_appearance = {}
    all_chars = set()
    for key in sorted(tracker.keys()):
        for name in tracker[key]:
            all_chars.add(name)
            if name not in first_appearance:
                first_appearance[name] = key
    # 입력된 캐릭터 목록과 비교
    input_chars = st.session_state.get("characters", "")
    warnings = []
    for name, unit in first_appearance.items():
        if unit != "01" and name not in input_chars:
            warnings.append(f"⚠️ '{name}' — UNIT {unit}에서 처음 등장. STEP 1 캐릭터 입력에 없는 인물.")
    return {
        "first_appearance": first_appearance,
        "warnings": warnings,
        "total": len(all_chars),
    }


def get_story_reinforcement_text() -> str:
    sr = st.session_state["story_reinforcement"]
    merged = []
    for k in ["기", "승", "전", "결"]:
        if sr.get(k):
            merged.append(f"[{k} 보강]\n{sr[k]}")
    merged_text = merge_nonempty(merged)
    st.session_state["story_reinforcement_merged"] = merged_text
    return merged_text


# ─────────────────────────────────────
# v3.13 M15-B: 자료 계량 수치 스캐너 (Unit 설계 재실행 없이 집필 단계에서 처리)
# ─────────────────────────────────────
# 스캔 대상 — (session_state 저장 위치, 표시 라벨)
#   컨셉 카드 4필드는 편집 가능(editable), 나머지는 진단만.
METRIC_SCAN_FIELDS = [
    ("f_overview", "작품 개요", True),
    ("f_characters", "캐릭터", True),
    ("f_synopsis", "줄거리/트리트먼트", True),
    ("f_notes", "추가 메모", True),
    ("f_locked_text", "고정 설정 (LOCKED)", True),
    ("f_open_text", "열린 설정 (OPEN)", True),
]


def _split_sentences_ko(text: str):
    """한국어 문장 대략 분리 — 스캔 리포트용."""
    import re as _re
    parts = _re.split(r"(?<=[.!?。])\s+|\n+", text)
    return [p.strip() for p in parts if p.strip()]


def scan_metric_expressions(max_items: int = 60) -> list:
    """v3.13 M15-B — 자료(컨셉 카드·설계안·보강본·Creator JSON)에서 계량 수치를 검출.

    Returns: [{"src_key","where","expr","sentence","editable"}, ...]
    본문(unit_drafts)은 스캔하지 않는다. 자료만 본다.
    """
    import re as _re
    pat_scan = _re.compile(METRIC_SCAN_RE)

    sources = []  # (src_key, where, text, editable)

    for key, label, editable in METRIC_SCAN_FIELDS:
        txt = st.session_state.get(key, "")
        if isinstance(txt, str) and txt.strip():
            sources.append((key, label, txt, editable))

    # 기승전결 보강본 (편집 가능)
    sr = st.session_state.get("story_reinforcement", {}) or {}
    for k in ["기", "승", "전", "결"]:
        txt = sr.get(k, "")
        if isinstance(txt, str) and txt.strip():
            sources.append((f"story_reinforcement::{k}", f"기승전결 보강 · {k}", txt, True))

    # Unit 설계안 (편집 가능 — 재생성 없이 문장만 교체)
    bp = st.session_state.get("unit_blueprints", {}) or {}
    for k in ["01-02", "03-04", "05-06", "07-08", "09-10", "11-12"]:
        txt = bp.get(k, "")
        if isinstance(txt, str) and txt.strip():
            sources.append((f"unit_blueprints::{k}", f"Unit {k} 설계", txt, True))

    # Creator 컨셉 카드 원본 JSON (읽기 전용 — 진단만)
    cj = st.session_state.get("creator_json_data")
    if isinstance(cj, (dict, list)):
        def _walk(o, path=""):
            if isinstance(o, dict):
                for kk, vv in o.items():
                    _walk(vv, f"{path}.{kk}" if path else str(kk))
            elif isinstance(o, list):
                for i, vv in enumerate(o):
                    _walk(vv, f"{path}[{i}]")
            elif isinstance(o, str) and o.strip():
                low = path.lower()
                if any(sk in low for sk in METRIC_SCAN_SKIP_KEYS):
                    return
                if pat_scan.search(o):
                    short = path.split(".")[-1] if path else "필드"
                    sources.append((None, f"Creator 카드 · {short}", o, False))
        _walk(cj)

    results = []
    seen = set()
    for src_key, where, text, editable in sources:
        for sent in _split_sentences_ko(text):
            exprs = [m.group(0) for m in pat_scan.finditer(sent)]
            if not exprs:
                continue
            for expr in dict.fromkeys(exprs):
                sig = (where, expr, sent[:40])
                if sig in seen:
                    continue
                seen.add(sig)
                results.append({
                    "src_key": src_key,
                    "where": where,
                    "expr": expr,
                    "sentence": sent,
                    "editable": editable,
                })
            if len(results) >= max_items:
                return results
    return results


def replace_in_source(src_key: str, old_sentence: str, new_sentence: str) -> bool:
    """v3.13 M15-B — 작가가 승인한 '한 문장'만 자료에서 교체한다.

    ★ 자동 일괄 치환은 하지 않는다. 작가가 직접 입력한 문장으로만 교체한다. ★
    Unit 본문(unit_drafts)은 이 함수의 대상이 아니다.
    """
    if not src_key or not old_sentence or old_sentence == new_sentence:
        return False

    if src_key.startswith("story_reinforcement::"):
        k = src_key.split("::", 1)[1]
        cur = st.session_state.get("story_reinforcement", {}).get(k, "")
        if old_sentence not in cur:
            return False
        st.session_state["story_reinforcement"][k] = cur.replace(old_sentence, new_sentence, 1)
        st.session_state["story_reinforcement_merged"] = ""  # 재조립 유도
        return True

    if src_key.startswith("unit_blueprints::"):
        k = src_key.split("::", 1)[1]
        cur = st.session_state.get("unit_blueprints", {}).get(k, "")
        if old_sentence not in cur:
            return False
        st.session_state["unit_blueprints"][k] = cur.replace(old_sentence, new_sentence, 1)
        return True

    cur = st.session_state.get(src_key, "")
    if not isinstance(cur, str) or old_sentence not in cur:
        return False
    st.session_state[src_key] = cur.replace(old_sentence, new_sentence, 1)
    return True


def gather_blueprints_text() -> str:
    bp = st.session_state["unit_blueprints"]
    keys = ["01-02", "03-04", "05-06", "07-08", "09-10", "11-12"]
    merged = []
    for key in keys:
        if bp.get(key):
            merged.append(f"[UNIT {key} 설계]\n{bp[key]}")
    return merge_nonempty(merged)


def gather_all_drafts_text() -> str:
    """전체 원고 — 내보내기용"""
    drafts = st.session_state["unit_drafts"]
    titles = st.session_state["chapter_titles"]
    merged = []
    for i in range(1, 14):
        key = f"{i:02d}" if i < 13 else "13"
        txt = drafts.get(key, "")
        if txt.strip():
            ch_title = titles.get(key, "")
            if ch_title:
                merged.append(f"{ch_title}\n{txt}")
            else:
                merged.append(txt)
    return merge_nonempty(merged)


def gather_chapter_titles_text() -> str:
    """현재 확정된 회차 제목을 목차 형태로 모은다. (v3.16)

    작가가 목차로 늘어놓았을 때의 인상을 그대로 검수 프롬프트에 넘긴다.
    """
    titles = st.session_state.get("chapter_titles", {}) or {}
    drafts = st.session_state.get("unit_drafts", {}) or {}
    lines = []
    for i in range(1, 14):
        key = f"{i:02d}" if i < 13 else "13"
        if not (drafts.get(key) or "").strip():
            continue
        t = (titles.get(key) or "").strip() or f"[CHAPTER {i}] (제목 없음)"
        lines.append(f"UNIT {i:02d}  {t}")
    return "\n".join(lines)


def gather_recent_drafts(current_unit: int, window: int = 2) -> str:
    """이전 Unit 요약 + 최근 N개 Unit 원고 — 연결성과 컨텍스트 동시 확보"""
    drafts = st.session_state["unit_drafts"]
    titles = st.session_state["chapter_titles"]
    summaries = st.session_state.get("unit_summaries", {})
    merged = []

    # ── 1. 이전 전체 Unit의 1줄 요약 (컨텍스트 유지) ──
    summary_lines = []
    for i in range(1, current_unit):
        key = f"{i:02d}" if i < 13 else "13"
        s = summaries.get(key, "")
        if s:
            summary_lines.append(f"  UNIT {key}: {s}")
    if summary_lines:
        merged.append("[이전 Unit 요약 — 전체 흐름 파악용]\n" + "\n".join(summary_lines))

    # ── 2. 최근 N개 Unit의 실제 텍스트 (연결성) ──
    start = max(1, current_unit - window)
    for i in range(start, current_unit):
        key = f"{i:02d}" if i < 13 else "13"
        txt = drafts.get(key, "")
        if txt.strip():
            ch_title = titles.get(key, "")
            label = ch_title if ch_title else f"[UNIT {key}]"
            if len(txt) > 3000:
                txt = "(...전략...)\n" + txt[-3000:]
            merged.append(f"{label}\n{txt}")
    return merge_nonempty(merged)


def export_txt(content: str) -> bytes:
    return export_clean_content(content).encode("utf-8")


def export_clean_content(content: str) -> str:
    """최종 원고에서 내부 마커를 제거하고 소설 포맷으로 정리"""
    import re
    lines = content.split("\n")
    cleaned = []
    for line in lines:
        s = line.strip()
        # 내부 라벨 제거
        if s.startswith("[UNIT ") and s.endswith("]"):
            continue
        # Stage A/B/C 내부 라벨 제거
        if s.startswith("# Chapter") and ("Stage A" in s or "Stage B" in s or "Stage C" in s):
            continue
        if s.startswith("# Chapter") and "Unit" in s:
            continue
        if s.startswith("## ") or s.startswith("### "):
            continue
        # markdown # CHAPTER → [CHAPTER] 변환
        m = re.match(r"^#\s*(CHAPTER\s*\d+)\s*[—\-]\s*(.+)$", s)
        if m:
            cleaned.append(f"[{m.group(1)}] — {m.group(2)}")
            continue
        # 기타 markdown # 헤더 → 일반 텍스트
        if s.startswith("# "):
            cleaned.append(s[2:])
            continue
        cleaned.append(line)

    result = "\n".join(cleaned)

    # ── 곧은 따옴표 → 둥근 따옴표 변환 ──
    # 큰따옴표: " → " / "
    result = re.sub(r'(?<=^)"', '\u201c', result, flags=re.MULTILINE)       # 줄 시작의 "
    result = re.sub(r'(?<=\s)"', '\u201c', result)                          # 공백 뒤의 "
    result = re.sub(r'"(?=\s)', '\u201d', result)                           # 공백 앞의 "
    result = re.sub(r'"(?=[.,!?\n])', '\u201d', result)                     # 구두점 앞의 "
    result = re.sub(r'"(?=$)', '\u201d', result, flags=re.MULTILINE)        # 줄 끝의 "
    # 남은 곧은 큰따옴표 처리 (짝 맞추기)
    remaining = []
    is_open = True
    for ch in result:
        if ch == '"':
            remaining.append('\u201c' if is_open else '\u201d')
            is_open = not is_open
        else:
            remaining.append(ch)
    result = "".join(remaining)

    # 작은따옴표: ' → ' / '  (대화 안 인용 등)
    result = result.replace("\u2018", "\u2018").replace("\u2019", "\u2019")  # 이미 둥근이면 유지
    # 곧은 작은따옴표는 한국 소설에서 거의 안 쓰이므로 최소 처리
    result = re.sub(r"(?<=\s)'", '\u2018', result)
    result = re.sub(r"'(?=[\s.,!?])", '\u2019', result)

    return result


def _looks_like_doc_title(line: str, title: str) -> bool:
    """첫 줄을 '작품 제목'으로 렌더해도 되는지 판정한다. (v3.15)

    제목이 아닌데 제목으로 처리하면 본문 첫 문장이 표지 제목 자리에
    16pt 굵게 중앙정렬로 박혀 나간다. 아래 조건을 모두 통과해야 제목이다.
      · 대사 부호로 시작하지 않는다 (대사는 제목이 될 수 없다)
      · 문장 종결부호로 끝나지 않는다
      · 40자 미만
      · 그리고 (전달받은 작품 제목과 일치) 또는 (제목 형태 라벨)
    """
    s = (line or "").strip()
    if not s or len(s) >= 40:
        return False
    if s[0] in '"\u201c\u300c\u2018\u201a\'\u300e(':
        return False
    if s[-1] in ".!?\u3002\uff01\uff1f\u201d\u300d":
        return False

    def _norm(t):
        return re.sub(r"[\s《》<>\[\]—·\-]+", "", (t or "")).lower()

    if title and _norm(s) and _norm(s) in _norm(title):
        return True
    # [CHAPTER n] / 제1장 같은 구조 라벨은 아래 전용 분기가 처리하므로 제외
    if s.startswith("[CHAPTER") or re.match(r"^제?\s*\d+\s*[장화부]", s):
        return False
    return False


def export_docx(title: str, content: str, spacing_mode: str = "standard") -> bytes:
    """한국 소설 원고 표준 DOCX — MS Word 소설 원고 포맷

    v3.16.2 — 빈 줄 처리 방식 개선.
    기존에는 본문의 빈 줄 하나하나를 '장면 전환'으로 보고 빈 단락을 넣었다.
    그런데 모델은 문단과 문단 사이에 항상 빈 줄을 넣으므로, 결과적으로
    전체 단락의 절반이 빈 단락이 됐다. (실측 — UNIT 01: 219단락 중 108개, 49%)
    줄간격 1.6까지 겹쳐 원고가 실제 분량의 두 배로 늘어났다.

    한국 소설 원고 표준은 문단 구분을 ★첫 줄 들여쓰기★로 한다.
    빈 줄은 장면 전환에만 쓴다. 그래서 기본값을 standard로 바꾼다.

    spacing_mode:
      "standard" — 빈 줄 1개는 문단 구분으로만 처리(빈 단락 없음).
                   빈 줄 2개 이상 또는 장면 전환 마커일 때만 여백을 준다.
      "relaxed"  — standard와 같되 문단 뒤에 6pt 여백을 준다.
      "web"      — 빈 줄을 그대로 빈 단락으로 옮긴다 (v3.16.1까지의 동작).
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    content = export_clean_content(content)

    doc = Document()

    # ── 페이지: A4, 여백 상하좌우 2.0cm (v3.16.3 — 작가 master 파일 규격) ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── 기본 스타일: 함초롱바탕 10.5pt, 줄간격 1.15, 양쪽정렬 (v3.16.3) ──
    # 작가가 master 파일에서 최종 스타일을 적용하므로, 엔진 출력은
    # master와 같은 규격으로 맞춰 붙여넣을 때 서식이 흔들리지 않게 한다.
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "HCR Batang"
    style_normal.font.size = Pt(10.5)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_before = Pt(0)
    # v3.16.2 — relaxed 모드는 빈 단락 대신 문단 여백으로 숨통을 준다.
    style_normal.paragraph_format.space_after = Pt(6) if spacing_mode == "relaxed" else Pt(0)
    style_normal.paragraph_format.first_line_indent = Cm(0.35)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 한글 폰트 (eastAsia)
    rpr = style_normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from lxml import etree
        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    # 한글 폰트 — 함초롱바탕 (v3.16.3)
    rfonts.set(qn("w:eastAsia"), "\ud568\ucd08\ub871\ubc14\ud0d5")
    rfonts.set(qn("w:ascii"), "HCR Batang")
    rfonts.set(qn("w:hAnsi"), "HCR Batang")

    # ── 헬퍼 ──
    def add_normal(text):
        return doc.add_paragraph(text)

    def add_dialogue(text):
        """대화문 — 들여쓰기 없음"""
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0)
        return p

    def add_centered(text, size=10.5, bold=False, before=0, after=0):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        if p.runs:
            p.runs[0].font.size = Pt(size)
            p.runs[0].font.bold = bold

    def add_scene_break():
        """장면 전환 — 빈 줄 1개"""
        p = doc.add_paragraph("")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)

    def add_page_break():
        """페이지 나누기"""
        from docx.oxml.ns import qn as _qn
        p = doc.add_paragraph()
        run = p.add_run()
        br = run._element.makeelement(_qn('w:br'), {_qn('w:type'): 'page'})
        run._element.append(br)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # ── 파싱 ──
    lines = content.split("\n")
    i = 0
    is_first_line = True
    chapter_count = 0

    while i < len(lines):
        s = lines[i].strip()

        # 빈 줄 처리 (v3.16.2)
        # 연속된 빈 줄의 개수를 세어, 문단 구분인지 장면 전환인지 가른다.
        if not s:
            blank_run = 0
            while i < len(lines) and not lines[i].strip():
                blank_run += 1
                i += 1
            if spacing_mode == "web":
                # 구버전 동작 — 빈 줄을 그대로 옮긴다
                add_scene_break()
            elif blank_run >= 2:
                # 빈 줄 2개 이상 = 명시적 장면 전환
                add_scene_break()
            # 빈 줄 1개 = 문단 구분. 들여쓰기가 이미 구분하므로 아무것도 넣지 않는다.
            continue

        # 장면 전환 마커 줄 — 빈 단락 하나로 치환 (v3.16.2)
        if re.fullmatch(r"[*·•\-—–\s]{1,12}", s) and not s.isdigit():
            add_scene_break()
            i += 1
            continue

        # 작품 제목 (첫 줄)
        # v3.15 수정 — 기존에는 '첫 줄이고 80자 미만'이면 무조건 작품 제목으로
        # 취급해 중앙정렬·굵게·16pt로 렌더했다. 그래서 모델이 [CHAPTER n] 헤더를
        # 빼먹은 Unit에서는 첫 대사("얼굴이 왜 그래.")가 작품 제목 자리에
        # 커다랗게 박혀 나갔다. (실제 사고 — UNIT 02 DOCX)
        # 이제 첫 줄이 실제로 제목일 때만 제목으로 처리한다.
        if is_first_line:
            is_first_line = False
            if _looks_like_doc_title(s, title):
                add_centered(s, size=16, bold=True, before=72, after=24)
                i += 1
                continue
            # 제목이 아니면 아래 일반 처리로 흘려보낸다.

        # 챕터 제목: [CHAPTER X] — ...
        if s.startswith("[CHAPTER"):
            if chapter_count > 0:
                add_page_break()
            add_centered(s, size=14, bold=True, before=36, after=18)
            chapter_count += 1
            i += 1
            continue

        # "끝."
        if s == "끝.":
            add_centered(s, size=11, bold=False, before=18, after=0)
            i += 1
            continue

        # 대화문 — 따옴표로 시작하면 들여쓰기 없음
        if s.startswith('"') or s.startswith('\u201c') or s.startswith('\u300c') or s.startswith("'"):
            add_dialogue(s)
            i += 1
            continue

        # 일반 본문 — 들여쓰기 적용 (기본 스타일)
        add_normal(s)
        i += 1

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def final_manuscript_text(current_title: str) -> str:
    parts = []
    if current_title.strip():
        parts.append(current_title.strip())

    # 본문
    drafts = gather_all_drafts_text().strip()
    if drafts:
        parts.append(drafts)
    manuscript = "\n\n".join(parts).rstrip()
    # v3.15.1 — 옛 세션에서 이어받은 Unit 본문에 중간 '끝.'이 남아 있으면
    # 통합본 한가운데에 그대로 박힌다. 내보내기 직전 한 번 더 걷어낸다.
    # 본문 문장은 건드리지 않는다 (작업 규칙 7).
    manuscript, _ = sanitize_manuscript(manuscript, is_final_unit=True)
    if manuscript and not manuscript.endswith("끝."):
        manuscript += "\n\n끝."
    return manuscript


def safe_filename(name: str) -> str:
    name = (name or "novel_draft").strip()
    name = re.sub(r"[^0-9A-Za-z가-힣_-]+", "_", name)
    return name or "novel_draft"


def parse_chapter_title(text: str) -> tuple:
    """원고 첫 줄에서 [CHAPTER X] — 서브타이틀을 추출한다.
    Returns (chapter_title, body) — 제목이 없으면 ("", text)"""
    if not text or not text.strip():
        return ("", text)
    lines = text.strip().split("\n", 1)
    first_line = lines[0].strip()
    if first_line.startswith("[CHAPTER"):
        body = lines[1].strip() if len(lines) > 1 else ""
        return (first_line, body)
    return ("", text.strip())


def set_status(message: str, status_type: str = "info") -> None:
    st.session_state["status_message"] = message
    st.session_state["status_type"] = status_type
    # v3.15 — 새 메시지는 아직 표시되지 않았음을 표시한다.
    st.session_state["status_shown"] = False


def clear_status() -> None:
    """상태 메시지를 즉시 지운다. (v3.15)"""
    st.session_state["status_message"] = ""
    st.session_state["status_type"] = "info"
    st.session_state["status_shown"] = False


def render_status() -> None:
    """상태 메시지를 표시한다.

    v3.15 — 기존에는 다음 set_status()가 호출될 때까지 메시지가 화면 상단에
    영구히 남았다. 특히 '결과가 중간에서 끊겼습니다' 오류가 한 번 뜨면
    이후 정상 작업 중에도 계속 붙어 있어 작가가 진행 중인 문제로 오해했다.
    이제 한 번 표시된 뒤에는 자동으로 소거된다.
    """
    msg = st.session_state.get("status_message", "").strip()
    status_type = st.session_state.get("status_type", "info")
    if not msg:
        return
    if st.session_state.get("status_shown"):
        # 이미 한 번 표시된 메시지 — 지우고 렌더하지 않는다.
        clear_status()
        return
    if status_type == "success":
        st.success(msg)
    elif status_type == "error":
        st.error(msg)
    elif status_type == "warning":
        st.warning(msg)
    else:
        st.info(msg)
    st.session_state["status_shown"] = True


def run_with_status(start_message: str, done_message: str, fn):
    set_status(start_message, "info")
    with st.spinner(start_message):
        try:
            result = fn()
            # v3.4.2 — 결과가 문장 중간에서 끊겼으면 완료로 처리하지 않는다.
            if isinstance(result, str) and looks_truncated(result):
                set_status(
                    "결과가 중간에서 끊겼습니다. 저장은 됐지만 끝부분을 확인하고 "
                    "필요하면 다시 생성하세요.",
                    "error",
                )
                st.warning(
                    "⚠️ **생성 결과가 문장 중간에서 끝났습니다.** "
                    f"마지막 부분: `…{result.rstrip()[-40:]}`  \n"
                    "토큰 상한에 걸렸을 가능성이 큽니다. 다시 생성하거나, "
                    "STEP 1의 입력 자료를 조금 줄여보세요."
                )
                return result
            set_status(done_message, "success")
            return result
        except Exception as e:
            set_status(f"작업 중 오류가 발생했습니다: {e}", "error")
            return None


def generate_or_expand_unit(unit_no: int, prompt: str) -> str:
    result = llm_call(prompt, max_tokens=MAX_TOKENS_LONG, use_opus=True)
    result = ensure_final_ending(result, unit_no)

    if is_incomplete_text(result, unit_no):
        expand_prompt = build_expand_incomplete_unit_prompt(
            unit_no=unit_no,
            current_text=result,
            target_length=UNIT_TARGET_LENGTHS.get(unit_no, 8000),
            min_length=UNIT_MIN_LENGTHS.get(unit_no, 6000),
        )
        extra = llm_call(expand_prompt, max_tokens=MAX_TOKENS_MID, use_opus=True)
        result = (result.rstrip() + "\n\n" + extra.strip()).strip()
        result = ensure_final_ending(result, unit_no)

    return result


# ─────────────────────────────────────
# HEADER
# ─────────────────────────────────────
st.markdown(
    f"""
<div class="header-wrap">
    <div class="header">BLUE JEANS PICTURES</div>
    <div class="brand-title">{APP_TITLE}</div>
    <div class="tagline">A <span>{APP_FORMAT}</span> ENGINE</div>
    <div class="tagline-desc">{APP_FORMAT_DESC}</div>
    <div class="sub">YOUNG · VINTAGE · FREE · INNOVATIVE</div>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="callout">
<b>시네마틱 노블</b>은 만화를 그래픽 노블이라 부르는 것과 같은 뜻의 형식입니다.
영화처럼 속도감 있게 전개되고 3막 15비트 구조를 갖추되, 영상 문법은 배제하고
소설만이 할 수 있는 묘사와 심리를 씁니다. 카메라가 볼 수 없는 것을 쓴다는 점이
시나리오와의 결정적 차이입니다.<br><br>
기획 자료를 넣으면 엔진이 분석 → 부족한 점 진단 → 기승전결 보강 → 12 Unit 설계 →
Unit 원고 생성 → 가제 검토/제목 제안까지 순서대로 진행합니다.
</div>
""",
    unsafe_allow_html=True,
)

render_status()

# ─────────────────────────────────────
# v3.1 SIDEBAR: 버전 및 활성 모듈 표시
# ─────────────────────────────────────
with st.sidebar:
    st.markdown(f"### 👖 Novel Engine {NOVEL_ENGINE_VERSION}")
    st.caption("Cinematic Novel Engine")
    st.caption(f"Build {NOVEL_ENGINE_BUILD_DATE}")
    st.caption(f"분석 {DEFAULT_MODEL} · 집필 {MODEL_OPUS}")

    # v3.4 작업 현황 + 저장 상기
    _sb_units = count_written_units()
    if _sb_units:
        st.markdown("---")
        st.markdown(f"**📝 Unit 원고 {_sb_units}개 작성됨**")
        st.caption("브라우저를 닫기 전에 상단 '프로젝트 저장'으로 백업하세요.")

    # 소설화 모드 상태 (출처별 구분)
    if st.session_state.get("scenario_fields_applied"):
        _sb_ex = st.session_state.get("scenario_extracted", {})
        _sb_src = _sb_ex.get("_source", "")
        if _sb_src == "creator_engine":
            st.success("🌱 Creator 기획 소설화 모드")
            st.caption(f"원작: {_sb_ex.get('_source_title', '')} ({_sb_ex.get('_source_format', '')})")
        elif _sb_src == "idea_engine":
            st.success("💡 Idea 기획 소설화 모드")
            st.caption(f"원작: {_sb_ex.get('_source_title', '')} ({_sb_ex.get('_source_format', '')})")
            _sb_pend = _sb_ex.get("_pending_items", [])
            if _sb_pend:
                st.caption(f"⚠️ 엔진 제안 {len(_sb_pend)}건 — 작가 확정 필요")
        else:
            st.success("📄 시나리오 소설화 모드 활성")
            stats = st.session_state.get("scenario_stats", {})
            st.caption(f"원작: {stats.get('char_count', 0):,}자 · {stats.get('scene_count', 0)}씬")

    st.markdown("---")
    st.markdown("**v3.1 신규**")
    st.markdown(
        f"""
- 📄 시나리오 업로드 모드
- 🧬 Sonnet 자동 추출 (STEP 1)
- 🗺️ 12 Unit 매핑 가이드 (STEP 4)
- 추출기: {'✅' if _SCENARIO_EXTRACTOR_AVAILABLE else '❌'}
"""
    )
    st.markdown("---")
    st.markdown("**v3.0 모듈**")
    st.markdown(
        f"""
- M1 BJND Scene Enforcer
- M2 OPENING MASTERY
- M3 BJND 4축 자기검증
- M4 Sub-genre OVERRIDE 4종
- M5 Profession Pack: {'✅' if _PROFESSION_PACK_AVAILABLE else '❌'}
- M6 Chapter Signature
- M7 Reader Retention Curve
- M8 POV Discipline
- M9 Period Pack: {'✅' if _PERIOD_PACK_AVAILABLE else '❌'}
- M10 Profession × Period 교차검증
"""
    )
    st.markdown("---")
    st.markdown("**BJND 임계치**")
    st.markdown(
        f"""
- 있었다 ≤ {BJND_THRESHOLDS['있었다']}회/Unit
- 것이었다 ≤ {BJND_THRESHOLDS['것이었다']}회/Unit
- 대사태그 ≤ {BJND_THRESHOLDS['대사태그']}회/Unit
- 현재형 ≤ {BJND_THRESHOLDS['현재형']}회/Unit (치명적)
- 계량수치 ≤ {BJND_THRESHOLDS['계량수치']}회/Unit (M15)
- 신체반응 ≤ {BJND_THRESHOLDS['신체반응']}회/Unit (v3.15)
- 서술자라벨링 ≤ {BJND_THRESHOLDS['기억오류']}회/Unit (M18)
- 회상마커 ≤ {BJND_THRESHOLDS['회상마커']}회/Unit (M12)
"""
    )
    st.caption(f"임계치 초과 시 자동 재생성 {AUTO_REGEN_MAX_RETRIES}회")

    st.markdown("**v3.15 신규**")
    _led_cnt = len([v for v in (st.session_state.get("continuity_ledger", {}) or {}).values() if v])
    st.markdown(
        f"""
- M19 설계안 준수 — Chapter 1 포함 전 경로 주입
- M17 연속성 원장 — {_led_cnt}개 Unit 기록됨
- M18 서술 거리 — 과잉노출·부당은폐 차단
"""
    )

# ─────────────────────────────────────
# v3.4 프로젝트 저장 / 불러오기
# ─────────────────────────────────────
_snap_units = count_written_units()
_snap_title = st.session_state.get("f_working_title", "").strip()
_has_work = bool(
    _snap_title
    or st.session_state.get("f_overview", "").strip()
    or _snap_units
    or st.session_state.get("scenario_fields_applied")
)

with st.expander(
    "💾 프로젝트 저장 / 불러오기"
    + (f"  —  {_snap_title or '제목 미정'} · Unit {_snap_units}개 작성됨" if _has_work else ""),
    expanded=not _has_work,
):
    st.caption(
        "작업 중인 모든 내용(STEP 1 자료·문체 분석·기승전결 보강·Unit 설계·Unit 원고)을 "
        "JSON 파일 하나로 저장하고 다시 불러옵니다. 브라우저를 닫아도 파일만 있으면 이어서 작업할 수 있습니다."
    )

    save_col, load_col = st.columns([1, 1])

    with save_col:
        st.markdown("**저장**")
        if _has_work:
            _payload = build_project_snapshot()
            _json_str = json.dumps(_payload, ensure_ascii=False, indent=2)
            _safe_title = re.sub(r'[\\\\/:*?"<>|]', "_", _snap_title or "novel")[:40]
            _fname = f"{_safe_title}_novelengine_{datetime.now().strftime('%Y%m%d_%H%M')}.json"
            st.download_button(
                "💾 프로젝트 저장 (JSON 다운로드)",
                data=_json_str.encode("utf-8"),
                file_name=_fname,
                mime="application/json",
                use_container_width=True,
                type="primary",
            )
            st.caption(
                f"저장 내용: STEP 1 자료 · Unit 원고 {_snap_units}개 · "
                f"설계/분석 결과 · 파일 크기 약 {len(_json_str) // 1024:,}KB"
            )
        else:
            st.button(
                "💾 프로젝트 저장 (JSON 다운로드)",
                use_container_width=True,
                disabled=True,
                help="저장할 작업 내용이 없습니다. STEP 1 입력 또는 원작 불러오기 후 사용하세요.",
            )
            st.caption("아직 저장할 내용이 없습니다.")

    with load_col:
        st.markdown("**불러오기**")
        _proj_file = st.file_uploader(
            "저장된 프로젝트 JSON 업로드",
            type=["json"],
            key="project_load_upload",
            label_visibility="collapsed",
        )
        if _proj_file is not None:
            _key = f"{_proj_file.name}_{_proj_file.size}"
            if st.session_state.get("_last_loaded_project") != _key:
                if st.button("📂 이 파일로 불러오기", use_container_width=True, type="primary"):
                    try:
                        _raw = _proj_file.read().decode("utf-8")
                        _payload_in = json.loads(_raw)
                    except Exception as e:
                        st.error(f"파일을 읽지 못했습니다: {e}")
                    else:
                        ok, msg, meta = restore_project_snapshot(_payload_in)
                        if not ok:
                            st.error(msg)
                        else:
                            st.session_state["_last_loaded_project"] = _key
                            st.success(
                                f"✅ 불러오기 완료 — {meta.get('title', '(제목 미정)')} "
                                f"(저장 시각 {meta.get('saved_at', '-')}, {msg})"
                            )
                            st.rerun()
            else:
                st.info("이미 불러온 파일입니다. 다른 파일을 선택하면 다시 불러올 수 있습니다.")
        st.caption("불러오면 현재 작업 내용은 덮어써집니다. 필요하면 먼저 저장하세요.")

# ─────────────────────────────────────
# v3.1 STEP 0 · 시나리오 업로드 (선택)
# ─────────────────────────────────────
st.markdown('<div class="section-header">📄 STEP 0 · 원작 불러오기 (선택 — 소설화 모드)</div>', unsafe_allow_html=True)

st.markdown(
    """
<div class="callout">
소설로 옮길 원작을 불러옵니다. <b>기존 원고</b>(시나리오 DOCX/TXT) · <b>Idea Engine JSON</b>(기획 씨앗) · <b>Creator Engine JSON</b>(확정 기획) 중 하나를 선택하세요.
어느 쪽이든 Sonnet이 STEP 1 입력 자료와 STEP 4 12 Unit 매핑 가이드를 자동 생성합니다.
추출 결과는 수정 가능한 형태로 STEP 1 필드에 자동 입력됩니다. 원작이 없으면 이 단계를 건너뛰고 STEP 1부터 수동 입력하세요.
</div>
""",
    unsafe_allow_html=True,
)

step0_tab_scenario, step0_tab_idea, step0_tab_creator = st.tabs(
    ["📄 기존 원고 (시나리오)", "💡 Idea Engine JSON", "🌱 Creator Engine JSON"]
)

with step0_tab_scenario:
    scenario_col1, scenario_col2 = st.columns([1, 1])

    with scenario_col1:
        scenario_file = st.file_uploader(
            "시나리오 파일 업로드 (.docx / .txt)",
            type=["docx", "txt"],
            key="scenario_upload",
            help="시나리오 원문 파일. Sonnet이 읽고 전체 구조를 추출합니다.",
        )

    with scenario_col2:
        scenario_pasted = st.text_area(
            "또는 시나리오 붙여넣기",
            height=120,
            placeholder="파일 업로드 대신 시나리오를 직접 붙여넣을 수 있습니다.",
            key="scenario_paste",
        )

    # 파일이 업로드되면 즉시 텍스트 추출
    if scenario_file is not None:
        try:
            file_bytes = scenario_file.read()
            if scenario_file.name.lower().endswith(".docx"):
                extracted_text = extract_text_from_docx(file_bytes)
            else:
                extracted_text = extract_text_from_txt(file_bytes)
            if extracted_text.strip():
                st.session_state["scenario_text"] = extracted_text
                st.session_state["scenario_stats"] = analyze_scenario_structure(extracted_text)
            else:
                st.error("파일에서 텍스트를 추출하지 못했습니다. TXT 인코딩을 확인하거나 붙여넣기를 사용해 주세요.")
        except Exception as e:
            st.error(f"파일 처리 오류: {e}")
    elif scenario_pasted and scenario_pasted.strip():
        st.session_state["scenario_text"] = scenario_pasted.strip()
        st.session_state["scenario_stats"] = analyze_scenario_structure(scenario_pasted.strip())

    # 업로드된 시나리오가 있으면 통계와 추출 버튼 표시
    if st.session_state.get("scenario_text", "").strip():
        stats = st.session_state.get("scenario_stats", {})
        st.markdown("**📊 시나리오 구조 통계**")
        stat_cols = st.columns(6)
        stat_cols[0].metric("글자 수", f"{stats.get('char_count', 0):,}")
        stat_cols[1].metric("문단 수", f"{stats.get('paragraph_count', 0):,}")
        stat_cols[2].metric("추정 씬 수", stats.get("scene_count", 0))
        stat_cols[3].metric("V.O 지시", stats.get("vo_count", 0))
        stat_cols[4].metric("CUT 지시", stats.get("cut_count", 0))
        stat_cols[5].metric("회상 씬", stats.get("flashback_count", 0))

        extract_col1, extract_col2 = st.columns([2, 1])
        with extract_col1:
            if st.button(
                "🧬 Sonnet 자동 추출 실행 — STEP 1 필드 + STEP 4 매핑 생성",
                type="primary",
                use_container_width=True,
                disabled=not _SCENARIO_EXTRACTOR_AVAILABLE,
            ):
                client = get_client()
                if client is None:
                    st.error("ANTHROPIC_API_KEY가 설정되지 않았습니다.")
                else:
                    with st.spinner("시나리오를 분석하고 있습니다... (30~60초 소요)"):
                        result = extract_scenario_fields(
                            scenario_text=st.session_state["scenario_text"],
                            anthropic_client=client,
                            model=DEFAULT_MODEL,
                            max_tokens=MAX_TOKENS_EXTRACT,
                        )
                    if "_error" in result:
                        st.error(f"추출 실패: {result['_error']}")
                        if result.get("_diagnostic"):
                            with st.expander("🔎 진단 정보", expanded=True):
                                for _dk, _dv in result["_diagnostic"].items():
                                    st.markdown(f"- **{_dk}**: {_dv}")
                        if "_raw_response" in result:
                            with st.expander("응답 원문 보기"):
                                st.text(result["_raw_response"])
                    else:
                        st.session_state["scenario_extracted"] = result
                        mapping_text = build_unit_mapping_text(result.get("unit_mapping", []))
                        st.session_state["scenario_mapping_text"] = mapping_text
                        st.session_state["scenario_fields_applied"] = True
                        apply_extracted_to_fields(result)
                        st.success(
                            "✅ 추출 완료. STEP 1 필드가 자동 입력되었고, STEP 4 Unit 설계 시 매핑 가이드가 자동 주입됩니다. "
                            "필요하면 STEP 1에서 수정하세요."
                        )
                        st.rerun()

        with extract_col2:
            if st.session_state.get("scenario_fields_applied"):
                if st.button("🔄 추출 결과 초기화", use_container_width=True):
                    st.session_state["scenario_extracted"] = {}
                    st.session_state["scenario_mapping_text"] = ""
                    st.session_state["scenario_fields_applied"] = False
                    st.rerun()

        # 추출 결과 미리보기
        extracted = st.session_state.get("scenario_extracted", {})
        if extracted and "_error" not in extracted:
            with st.expander("🔍 추출 결과 미리보기", expanded=False):
                st.markdown(f"**로그라인:** {extracted.get('logline', '')}")
                st.markdown(f"**장르:** {extracted.get('genre', '')}")
                st.markdown(f"**주인공 직업:** {extracted.get('profession_protagonist', '')}")
                st.markdown(f"**적대자/조연 직업:** {extracted.get('profession_antagonist', '')}")
                st.markdown(f"**시대 키:** {extracted.get('period_keys', [])}")
                st.markdown("**작품 개요:**")
                st.text(extracted.get("overview", ""))
                st.markdown("**캐릭터:**")
                st.text(extracted.get("characters", "")[:1000] + ("..." if len(extracted.get("characters", "")) > 1000 else ""))
                st.markdown("**12 Unit 매핑 가이드:**")
                mapping = extracted.get("unit_mapping", [])
                for item in mapping:
                    st.markdown(f"- **Unit {item.get('unit_no', '?')}**: {item.get('function', '')}")

with step0_tab_idea:
    st.markdown(
        "**Idea Engine이 뽑은 기획 씨앗(IdeaSeed) JSON을 업로드하면, "
        "기획 단계 자료를 소설 언어로 번역해 STEP 1 필드와 STEP 4 매핑을 자동 생성합니다.**"
    )
    st.caption(
        "Idea 단계는 아직 결정되지 않은 항목이 남아 있습니다. 엔진이 각 미결정에 대해 제안을 채워 넣되, "
        "해당 자리에 [엔진 제안 — 작가 확정 필요] 표식을 붙입니다. 무엇이 원본 확정이고 무엇이 엔진 판단인지 "
        "STEP 1에서 바로 구분할 수 있습니다."
    )

    idea_file = st.file_uploader(
        "Idea Engine JSON 업로드 (.json)",
        type=["json"],
        key="idea_json_upload",
        help="Idea Engine v2.0+ 이 저장한 IdeaSeed JSON 파일",
        disabled=not _IDEA_EXTRACTOR_AVAILABLE,
    )

    if idea_file is not None:
        try:
            _idata = load_idea_json(idea_file.read())
            if "_error" in _idata:
                st.error(f"JSON 로드 실패: {_idata['_error']}")
            elif not is_idea_json(_idata):
                st.warning(
                    "이 파일은 Idea Engine 출력으로 보이지 않습니다. "
                    "Creator Engine JSON이라면 오른쪽 'Creator Engine JSON' 탭을 사용하세요."
                )
                st.session_state["idea_json_data"] = {}
                st.session_state["idea_json_meta"] = {}
                st.session_state["idea_pending_items"] = []
            else:
                st.session_state["idea_json_data"] = _idata
                st.session_state["idea_json_meta"] = get_idea_meta(_idata)
                st.session_state["idea_pending_items"] = collect_pending_items(_idata)
        except Exception as e:
            st.error(f"파일 처리 오류: {e}")

    _imeta = st.session_state.get("idea_json_meta", {})
    if _imeta:
        st.markdown("**📋 Idea 기획 정보**")
        im_cols = st.columns(4)
        im_cols[0].metric("제목", _imeta.get("title", "-"))
        im_cols[1].metric("원본 매체", _imeta.get("format", "-"))
        im_cols[2].metric("판정", _imeta.get("verdict", "-"))
        im_cols[3].metric("미결정", f"{_imeta.get('pending_count', '0')}건")
        st.caption(
            f"엔진: {_imeta.get('engine', '-')} {_imeta.get('engine_version', '')} · "
            f"장르: {_imeta.get('genre', '-')} · Hook Score: {_imeta.get('hook_score', '-')}"
        )

        # 미결정 항목 표시 — 작가가 STEP 1에서 우선 검토할 지점
        _pend = st.session_state.get("idea_pending_items", [])
        if _pend:
            with st.expander(f"⚠️ 원본 미결정 {len(_pend)}건 — 엔진이 제안으로 채웁니다", expanded=True):
                st.caption(
                    "아래 항목은 Idea 단계에서 확정되지 않았습니다. 변환 시 엔진이 각각 제안을 선택해 "
                    "[엔진 제안 — 작가 확정 필요] 표식과 함께 STEP 1에 넣습니다. 작가가 검토 후 확정하세요."
                )
                for _i, _p in enumerate(_pend, 1):
                    _imp = _p.get("importance", "")
                    _badge = f" `{_imp}`" if _imp else ""
                    st.markdown(f"**{_i}. {_p.get('question', '')}**{_badge}")
                    for _o in _p.get("options", []):
                        st.markdown(f"    - {_o}")

        iv_col1, iv_col2 = st.columns([2, 1])
        with iv_col1:
            if st.button(
                "💡 Idea JSON → 소설화 변환 (미결정은 엔진 제안으로 채움)",
                type="primary",
                use_container_width=True,
                key="convert_idea_btn",
                disabled=not _IDEA_EXTRACTOR_AVAILABLE,
            ):
                client = get_client()
                if client is None:
                    st.error("ANTHROPIC_API_KEY가 설정되지 않았습니다.")
                else:
                    with st.spinner("기획 씨앗을 소설로 번역하고 미결정 항목을 제안하고 있습니다... (30~60초 소요)"):
                        result = extract_idea_fields(
                            idea_json=st.session_state["idea_json_data"],
                            anthropic_client=client,
                            model=DEFAULT_MODEL,
                            max_tokens=MAX_TOKENS_EXTRACT,
                        )
                    if "_error" in result:
                        st.error(f"변환 실패: {result['_error']}")
                        if result.get("_diagnostic"):
                            with st.expander("🔎 진단 정보", expanded=True):
                                for _dk, _dv in result["_diagnostic"].items():
                                    st.markdown(f"- **{_dk}**: {_dv}")
                        if "_raw_response" in result:
                            with st.expander("응답 원문 보기"):
                                st.text(result["_raw_response"])
                    else:
                        st.session_state["scenario_extracted"] = result
                        mapping_text = build_unit_mapping_text(result.get("unit_mapping", []))
                        st.session_state["scenario_mapping_text"] = mapping_text
                        st.session_state["scenario_fields_applied"] = True
                        apply_extracted_to_fields(result)
                        st.success(
                            "✅ 소설화 변환 완료. STEP 1 필드가 자동 입력되었습니다. "
                            "[엔진 제안 — 작가 확정 필요] 표식이 붙은 항목을 우선 검토하세요."
                        )
                        st.rerun()

        with iv_col2:
            if st.session_state.get("scenario_fields_applied"):
                if st.button("🔄 변환 결과 초기화", use_container_width=True, key="reset_idea_btn"):
                    st.session_state["scenario_extracted"] = {}
                    st.session_state["scenario_mapping_text"] = ""
                    st.session_state["scenario_fields_applied"] = False
                    st.rerun()

        # 변환 결과 미리보기 (Idea 출처인 경우만)
        _ires = st.session_state.get("scenario_extracted", {})
        if _ires and "_error" not in _ires and _ires.get("_source") == "idea_engine":
            with st.expander("🔍 소설화 변환 결과 미리보기", expanded=False):
                st.markdown(f"**로그라인:** {_ires.get('logline', '')}")
                st.markdown(f"**장르(소설):** {_ires.get('genre', '')}")
                st.markdown(f"**주인공 직업:** {_ires.get('profession_protagonist', '')}")
                st.markdown(f"**적대자/조연 직업:** {_ires.get('profession_antagonist', '')}")
                st.markdown(f"**시대 키:** {_ires.get('period_keys', [])}")
                st.markdown("**작품 개요:**")
                st.text(_ires.get("overview", ""))
                st.markdown("**캐릭터:**")
                _ichtxt = _ires.get("characters", "")
                st.text(_ichtxt[:1200] + ("..." if len(_ichtxt) > 1200 else ""))
                st.markdown("**12 Unit 매핑 가이드:**")
                for item in _ires.get("unit_mapping", []):
                    st.markdown(f"- **Unit {item.get('unit_no', '?')}**: {item.get('function', '')}")

with step0_tab_creator:
    st.markdown(
        "**Creator Engine이 뽑은 기획 JSON(영화·시리즈 format)을 업로드하면, "
        "영상 서사를 소설 언어로 번역해 STEP 1 필드와 STEP 4 매핑을 자동 생성합니다.**"
    )
    st.caption(
        "Creator Engine의 백업 JSON 파일을 넣으세요. 캐릭터 바이블·세계관·3막 구조·톤 재료까지 "
        "최대한 보존해 소설화합니다. 영상 전용 연출(씬·컷·지문)은 소설 서술로 자동 전환됩니다."
    )

    creator_file = st.file_uploader(
        "Creator Engine JSON 업로드 (.json)",
        type=["json"],
        key="creator_json_upload",
        help="Creator Engine v2.5+ 이 저장한 기획 JSON 백업 파일",
        disabled=not _CREATOR_EXTRACTOR_AVAILABLE,
    )

    if creator_file is not None:
        try:
            data = load_creator_json(creator_file.read())
            if "_error" in data:
                st.error(f"JSON 로드 실패: {data['_error']}")
            elif not is_creator_json(data):
                st.warning(
                    "이 파일은 Creator Engine 출력으로 보이지 않습니다. "
                    "Idea Engine 출력이나 다른 엔진의 세이브 파일일 수 있습니다. "
                    "Creator Engine의 기획 JSON을 확인해 주세요."
                )
                st.session_state["creator_json_data"] = {}
                st.session_state["creator_json_meta"] = {}
            else:
                st.session_state["creator_json_data"] = data
                st.session_state["creator_json_meta"] = get_creator_meta(data)
        except Exception as e:
            st.error(f"파일 처리 오류: {e}")

    # 로드된 Creator JSON이 있으면 메타 카드 + 변환 버튼
    _cmeta = st.session_state.get("creator_json_meta", {})
    if _cmeta:
        st.markdown("**📋 Creator 기획 정보**")
        cm_cols = st.columns(4)
        cm_cols[0].metric("제목", _cmeta.get("title", "-"))
        cm_cols[1].metric("원본 매체", _cmeta.get("format", "-"))
        cm_cols[2].metric("장르", _cmeta.get("genre", "-"))
        cm_cols[3].metric("캐릭터", f"{_cmeta.get('char_count', '0')}명")
        st.caption(
            f"엔진: {_cmeta.get('engine', '-')} {_cmeta.get('engine_version', '')} · "
            f"단계: {_cmeta.get('stage', '-')} · 완성도: {_cmeta.get('final_score', '-')}"
        )

        cv_col1, cv_col2 = st.columns([2, 1])
        with cv_col1:
            if st.button(
                "🌱 Creator JSON → 소설화 변환 (STEP 1 필드 + STEP 4 매핑 생성)",
                type="primary",
                use_container_width=True,
                key="convert_creator_btn",
                disabled=not _CREATOR_EXTRACTOR_AVAILABLE,
            ):
                client = get_client()
                if client is None:
                    st.error("ANTHROPIC_API_KEY가 설정되지 않았습니다.")
                else:
                    with st.spinner("영상 기획을 소설로 번역하고 있습니다... (30~60초 소요)"):
                        result = extract_creator_fields(
                            creator_json=st.session_state["creator_json_data"],
                            anthropic_client=client,
                            model=DEFAULT_MODEL,
                            max_tokens=MAX_TOKENS_EXTRACT,
                        )
                    if "_error" in result:
                        st.error(f"변환 실패: {result['_error']}")
                        if result.get("_diagnostic"):
                            with st.expander("🔎 진단 정보", expanded=True):
                                for _dk, _dv in result["_diagnostic"].items():
                                    st.markdown(f"- **{_dk}**: {_dv}")
                        if "_raw_response" in result:
                            with st.expander("응답 원문 보기"):
                                st.text(result["_raw_response"])
                    else:
                        st.session_state["scenario_extracted"] = result
                        mapping_text = build_unit_mapping_text(result.get("unit_mapping", []))
                        st.session_state["scenario_mapping_text"] = mapping_text
                        st.session_state["scenario_fields_applied"] = True
                        apply_extracted_to_fields(result)
                        st.success(
                            "✅ 소설화 변환 완료. STEP 1 필드가 자동 입력되었고, "
                            "STEP 4 Unit 설계 시 매핑 가이드가 자동 주입됩니다. "
                            "필요하면 STEP 1에서 수정하세요."
                        )
                        st.rerun()

        with cv_col2:
            if st.session_state.get("scenario_fields_applied"):
                if st.button("🔄 변환 결과 초기화", use_container_width=True, key="reset_creator_btn"):
                    st.session_state["scenario_extracted"] = {}
                    st.session_state["scenario_mapping_text"] = ""
                    st.session_state["scenario_fields_applied"] = False
                    st.rerun()

        # 변환 결과 미리보기 (Creator 출처인 경우만)
        _cres = st.session_state.get("scenario_extracted", {})
        if _cres and "_error" not in _cres and _cres.get("_source") == "creator_engine":
            with st.expander("🔍 소설화 변환 결과 미리보기", expanded=False):
                st.markdown(f"**로그라인:** {_cres.get('logline', '')}")
                st.markdown(f"**장르(소설):** {_cres.get('genre', '')}")
                st.markdown(f"**주인공 직업:** {_cres.get('profession_protagonist', '')}")
                st.markdown(f"**적대자/조연 직업:** {_cres.get('profession_antagonist', '')}")
                st.markdown(f"**시대 키:** {_cres.get('period_keys', [])}")
                st.markdown("**작품 개요:**")
                st.text(_cres.get("overview", ""))
                st.markdown("**캐릭터:**")
                _chtxt = _cres.get("characters", "")
                st.text(_chtxt[:1200] + ("..." if len(_chtxt) > 1200 else ""))
                st.markdown("**12 Unit 매핑 가이드:**")
                for item in _cres.get("unit_mapping", []):
                    st.markdown(f"- **Unit {item.get('unit_no', '?')}**: {item.get('function', '')}")


# v3.1: STEP 1 필드 기본값 결정 (추출 결과 있으면 그걸 사용)
_ex = st.session_state.get("scenario_extracted", {}) if st.session_state.get("scenario_fields_applied") else {}

# ─────────────────────────────────────
# STEP 1
# ─────────────────────────────────────
st.markdown('<div class="section-header">🔥 STEP 1 · 작품 자료 입력</div>', unsafe_allow_html=True)

if st.session_state.get("scenario_fields_applied"):
    _src = _ex.get("_source", "")
    _src_title = _ex.get("_source_title", "")
    _src_fmt = _ex.get("_source_format", "")
    if _src == "creator_engine":
        st.info(
            f"🌱 Creator 기획 '{_src_title}'({_src_fmt})을 소설화한 결과가 자동 입력되었습니다. "
            "필요하면 아래 필드를 수정하세요."
        )
    elif _src == "idea_engine":
        _pend_list = _ex.get("_pending_items", [])
        st.info(
            f"💡 Idea 기획 '{_src_title}'({_src_fmt})을 소설화한 결과가 자동 입력되었습니다. "
            "필요하면 아래 필드를 수정하세요."
        )
        if _pend_list:
            st.warning(
                f"⚠️ 원본 미결정 {len(_pend_list)}건은 엔진 제안으로 채워졌습니다. "
                "아래 필드에서 **[엔진 제안 — 작가 확정 필요]** 표식이 붙은 부분을 우선 검토하세요."
            )
            with st.expander("미결정이었던 항목 보기", expanded=False):
                for _i, _q in enumerate(_pend_list, 1):
                    st.markdown(f"{_i}. {_q}")
    else:
        st.info("📄 시나리오 추출 결과가 자동 입력되었습니다. 필요하면 아래 필드를 수정하세요.")

col1, col2 = st.columns([1, 1])

with col1:
    working_title = st.text_input(
        "현재 가제",
        key="f_working_title",
        placeholder="예: 감각구역 / 머지 앤 어퀴지션 / 검은 항구",
    )
    genre = st.text_input(
        "장르",
        key="f_genre",
        placeholder="예: 스릴러, 역사드라마, 금융 스릴러, 첩보물",
    )
    format_mode = st.selectbox(
        "형식", ["장편소설", "웹소설", "하이브리드"], key="f_format_mode"
    )

with col2:
    pov = st.selectbox(
        "시점", ["3인칭 제한", "1인칭", "듀얼 POV", "다중시점"], key="f_pov"
    )
    target_length = st.text_input(
        "목표 분량", key="f_target_length",
        placeholder="예: 12만자 / 12 Units / Unit당 1만자",
    )
    style_strength = st.selectbox(
        "문체 반영 강도", ["약", "중", "강"], key="f_style_strength"
    )

overview = st.text_area(
    "작품 개요",
    height=220,
    key="f_overview",
    placeholder="로그라인, 기획의도, 세계관, 장르 톤, 작품의 핵심 질문, 차별점",
)

characters = st.text_area(
    "캐릭터",
    height=220,
    key="f_characters",
    placeholder="주인공 / 적대자 / 조력자 / 핵심 관계, 각 인물의 욕망 / 결핍 / 비밀 / 변화",
)

synopsis = st.text_area(
    "줄거리 / 트리트먼트",
    height=260,
    key="f_synopsis",
    placeholder="시작, 중반, 위기, 클라이맥스, 엔딩 방향, 반드시 살릴 사건",
)

notes = st.text_area(
    "추가 메모 (선택)",
    height=180,
    key="f_notes",
    placeholder="약한 부분, 반드시 살릴 장면, 정보 레이어, 역사 고증 메모, 참고 톤",
)

style_sample = st.text_area(
    "문체 샘플 (선택)",
    height=220,
    key="f_style_sample",
    placeholder="Mr.MOON이 직접 쓴 소설/산문/블로그 문장 일부",
)

# ─────────────────────────────────────
# v3.0 신규: 직업(M5) + 시대(M9) 입력 섹션
# ─────────────────────────────────────
st.markdown(
    '<div style="margin-top:16px; font-weight:600; color:#191970;">🎯 v3.0 · 전문성 강화 (직업 / 시대)</div>',
    unsafe_allow_html=True,
)
st.caption(
    "입력하신 정보로 Creator Engine의 Profession Pack(19 카테고리) 및 Period Pack(10 시대)이 자동 주입됩니다. "
    "비워두면 주입하지 않습니다."
)

prof_col1, prof_col2 = st.columns([1, 1])

with prof_col1:
    profession_protagonist = st.text_input(
        "주인공 직업 (M5)",
        key="f_profession_protagonist",
        placeholder="예: M&A 변호사 / 강력계 형사 / 오너 셰프 / 투자은행 VP / 군의관",
        help="Profession Pack이 자동 감지하여 전문 용어·공간·일상·스트레스를 주입합니다.",
    )

with prof_col2:
    profession_antagonist = st.text_input(
        "주요 조연/적대자 직업 (M5, 선택)",
        key="f_profession_antagonist",
        placeholder="예: 검사 / 조직폭력 보스 / 기자 / 로비스트",
        help="주인공과 다른 직업이면 추가 주입. 같으면 중복 방지.",
    )

# v3.0 M9: 시대 설정
period_col1, period_col2 = st.columns([1, 2])

with period_col1:
    period_mode = st.radio(
        "시대 모드 (M9)",
        ["현대 (시대 주입 없음)", "자동 감지 (LOCKED에서)", "수동 선택"],
        key="f_period_mode",
        help="역사소설이 아니면 '현대'를 유지하세요. 자동 감지는 LOCKED 블록의 연도·인물·사건 키워드를 스캔합니다.",
    )

with period_col2:
    period_keys_selected = []
    if period_mode == "수동 선택" and PERIOD_KEYS:
        # 한국어 라벨로 표시하되 내부 키로 저장
        period_options = [f"{k} · {get_period_label(k)}" for k in PERIOD_KEYS]
        # 저장된 라벨 중 현재 옵션에 존재하는 것만 유지 (불러오기 안전)
        _saved_labels = [
            x for x in st.session_state.get("f_period_labels", [])
            if x in period_options
        ]
        if _saved_labels != st.session_state.get("f_period_labels", []):
            st.session_state["f_period_labels"] = _saved_labels
        selected_labels = st.multiselect(
            "시대 선택 (최대 2개, 다중 시대 교차 전개 시 2개)",
            period_options,
            key="f_period_labels",
            max_selections=2,
            help="일제강점기 + 현대, 구한말 + 일제강점기 등 교차 전개도 가능.",
        )
        # label에서 키만 추출
        period_keys_selected = [x.split(" · ")[0] for x in selected_labels]
    elif period_mode == "자동 감지 (LOCKED에서)":
        st.caption("ℹ️ LOCKED 블록 입력 후 Unit 생성 시 자동 감지됩니다.")

lock_col1, lock_col2 = st.columns([1, 1])

with lock_col1:
    locked_text = st.text_area(
        "🔒 LOCKED 설정 (절대 변경 불가)",
        height=180,
        key="f_locked_text",
        placeholder="변경 금지 항목을 줄 단위로 입력\n예:\n- 한유진: QLCP 대표. 직책 변경 금지.\n- 마이클 모건: 적대자. 동맹으로 변경 금지.\n- 기획의도: 글로벌 금융 권력 비판이 테마에 반영되어야 함.",
    )

with lock_col2:
    open_text = st.text_area(
        "🔓 OPEN 설정 (창작 가능 범위)",
        height=180,
        key="f_open_text",
        placeholder="자유롭게 창작 가능한 항목\n예:\n- 캐릭터 외형, 습관, 말투 디테일은 자유롭게 확장 가능.\n- 장면별 감정 변화와 감각 묘사는 자유롭게 창작 가능.",
    )

locked_block = build_locked_block(locked_text, open_text)

# v3.0: 직업/시대 정보를 하나의 문자열과 키 리스트로 정리
profession_text_combined = " / ".join(
    p.strip() for p in [profession_protagonist, profession_antagonist] if p and p.strip()
)

# 시대 키 확정
if period_mode == "수동 선택":
    active_period_keys = period_keys_selected
elif period_mode == "자동 감지 (LOCKED에서)":
    # 자동 감지는 각 생성 함수 내부에서 locked_block 기반으로 수행
    active_period_keys = None
else:
    active_period_keys = []  # 빈 리스트 = 주입 안 함

# 감지 미리보기 (자동 감지 모드일 때만)
if period_mode == "자동 감지 (LOCKED에서)" and locked_text.strip() and _PERIOD_PACK_AVAILABLE:
    try:
        preview = detect_period_from_locked(locked_text)
        if preview:
            preview_labels = [get_period_label(k) for k in preview[:2]]
            st.info(f"🕰️ 감지된 시대: {' · '.join(preview_labels)}")
        else:
            st.caption("ℹ️ LOCKED에서 시대 키워드를 감지하지 못했습니다. 현대로 처리됩니다.")
    except Exception:
        pass

# ─────────────────────────────────────
# STEP 2
# ─────────────────────────────────────
st.markdown('<div class="section-header">🔬 STEP 2 · 문체 / 분석</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="small-meta">2-1은 문체 샘플이 있을 때만. 2-2 → 2-3 순서로 실행합니다.</div>',
    unsafe_allow_html=True,
)

c1, c2, c3 = st.columns(3)

with c1:
    if st.button("2-1 · 문체 샘플 분석", type="primary", use_container_width=True):
        def _job():
            prompt = STYLE_DNA_ANALYSIS_PROMPT.format(style_sample=style_sample or "샘플 없음")
            return llm_call(prompt, max_tokens=MAX_TOKENS_ANALYSIS)
        result = run_with_status("문체 샘플을 분석 중입니다...", "문체 샘플 분석이 완료되었습니다.", _job)
        if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
            st.session_state["style_dna"] = result

with c2:
    if st.button("2-2 · 기획서 통합 분석", use_container_width=True):
        def _job():
            prompt = build_merge_analysis_prompt(
                working_title=working_title,
                genre=genre,
                format_mode=format_mode,
                pov=pov,
                target_length=target_length,
                overview=overview,
                characters=characters,
                synopsis=synopsis,
                notes=notes,
                style_dna=st.session_state["style_dna"],
                style_strength=style_strength,
                locked_block=locked_block,
            )
            return llm_call(prompt, max_tokens=MAX_TOKENS_ANALYSIS)
        result = run_with_status("기획서 통합 분석 중입니다...", "기획서 통합 분석이 완료되었습니다.", _job)
        if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
            st.session_state["merged_analysis"] = result

with c3:
    if st.button("2-3 · 부족한 점 진단", use_container_width=True):
        def _job():
            prompt = build_gap_diagnosis_prompt(
                working_title=working_title,
                merged_analysis=st.session_state["merged_analysis"],
                overview=overview,
                characters=characters,
                synopsis=synopsis,
                notes=notes,
                style_dna=st.session_state["style_dna"],
                locked_block=locked_block,
            )
            return llm_call(prompt, max_tokens=MAX_TOKENS_ANALYSIS)
        result = run_with_status("부족한 점을 진단 중입니다...", "부족한 점 진단이 완료되었습니다.", _job)
        if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
            st.session_state["gap_diagnosis"] = result

if st.session_state["style_dna"]:
    with st.expander("STYLE DNA 보기", expanded=False):
        st.markdown(st.session_state["style_dna"])

if st.session_state["merged_analysis"]:
    with st.expander("기획서 통합 분석 보기", expanded=False):
        st.markdown(st.session_state["merged_analysis"])

if st.session_state["gap_diagnosis"]:
    with st.expander("부족한 점 진단 보기", expanded=False):
        st.markdown(st.session_state["gap_diagnosis"])

# ─────────────────────────────────────
# STEP 3
# ─────────────────────────────────────
st.markdown('<div class="section-header">📖 STEP 3 · 전체 줄거리 보강 (기승전결 분할)</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="small-meta">3-1 기 → 3-2 승 → 3-3 전 → 3-4 결 순서로 실행합니다. '
    '앞 단계 결과를 뒤 단계가 받아 씁니다.</div>',
    unsafe_allow_html=True,
)

sr_col1, sr_col2, sr_col3, sr_col4 = st.columns(4)

def reinforce_segment(segment_name: str):
    def _job():
        prompt = build_story_reinforcement_prompt(
            segment_name=segment_name,
            working_title=working_title,
            genre=genre,
            overview=overview,
            characters=characters,
            synopsis=synopsis,
            notes=notes,
            merged_analysis=st.session_state["merged_analysis"],
            gap_diagnosis=st.session_state["gap_diagnosis"],
            style_dna=st.session_state["style_dna"],
            locked_block=locked_block,
        )
        return llm_call(prompt, max_tokens=MAX_TOKENS_DESIGN)

    result = run_with_status(
        f"{segment_name} 구간을 장편소설 구조로 보강 중입니다...",
        f"{segment_name} 구간 보강이 완료되었습니다.",
        _job,
    )
    if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
        st.session_state["story_reinforcement"][segment_name] = result
        get_story_reinforcement_text()

with sr_col1:
    if st.button("3-1 · 기 보강", use_container_width=True):
        reinforce_segment("기")
with sr_col2:
    if st.button("3-2 · 승 보강", use_container_width=True):
        reinforce_segment("승")
with sr_col3:
    if st.button("3-3 · 전 보강", use_container_width=True):
        reinforce_segment("전")
with sr_col4:
    if st.button("3-4 · 결 보강", use_container_width=True):
        reinforce_segment("결")

story_merged_text = get_story_reinforcement_text()

for seg in ["기", "승", "전", "결"]:
    seg_text = st.session_state["story_reinforcement"].get(seg, "")
    if seg_text:
        with st.expander(f"{seg} 보강 보기", expanded=False):
            st.markdown(seg_text)

# ─────────────────────────────────────
# STEP 4
# ─────────────────────────────────────
st.markdown('<div class="section-header">🏗️ STEP 4 · 12 Unit 설계 (2 Unit씩 6개 버튼)</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="small-meta">4-1부터 순서대로 실행하세요. '
    '앞 그룹 설계가 뒤 그룹의 전제가 되므로 건너뛰면 연결이 어긋납니다.</div>',
    unsafe_allow_html=True,
)

bp_cols_top = st.columns(3)
bp_cols_bottom = st.columns(3)

def build_blueprint(group_key: str):
    def _job():
        prompt = build_unit_blueprint_prompt(
            group_key=group_key,
            working_title=working_title,
            genre=genre,
            format_mode=format_mode,
            pov=pov,
            overview=overview,
            characters=characters,
            story_reinforcement_merged=story_merged_text,
            synopsis=synopsis,
            notes=notes,
            style_dna=st.session_state["style_dna"],
            locked_block=locked_block,
            # v3.0 신규
            profession_text=profession_text_combined,
            period_keys=active_period_keys,
            # v3.1 신규: 시나리오 소설화 매핑 가이드
            scenario_mapping=st.session_state.get("scenario_mapping_text", ""),
            # v3.14 M16
            food_signature=st.session_state.get("signature_food_opening", True),
        )
        return llm_call(prompt, max_tokens=MAX_TOKENS_DESIGN)

    result = run_with_status(
        f"UNIT {group_key} 설계를 생성 중입니다...",
        f"UNIT {group_key} 설계가 완료되었습니다.",
        _job,
    )
    if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
        st.session_state["unit_blueprints"][group_key] = result

buttons = [
    ("4-1 · UNIT 01~02 설계", "01-02"),
    ("4-2 · UNIT 03~04 설계", "03-04"),
    ("4-3 · UNIT 05~06 설계", "05-06"),
    ("4-4 · UNIT 07~08 설계", "07-08"),
    ("4-5 · UNIT 09~10 설계", "09-10"),
    ("4-6 · UNIT 11~12 설계", "11-12"),
]

for idx, (label, group_key) in enumerate(buttons):
    target_col = bp_cols_top[idx] if idx < 3 else bp_cols_bottom[idx - 3]
    with target_col:
        if st.button(label, use_container_width=True):
            build_blueprint(group_key)

for group_key in ["01-02", "03-04", "05-06", "07-08", "09-10", "11-12"]:
    if st.session_state["unit_blueprints"].get(group_key):
        with st.expander(f"UNIT {group_key} 설계 보기", expanded=False):
            st.markdown(st.session_state["unit_blueprints"][group_key])

all_blueprints_text = gather_blueprints_text()

# ─────────────────────────────────────
# STEP 5
# ─────────────────────────────────────
st.markdown('<div class="section-header">✍️ STEP 5 · Unit 원고 생성 / 다시 쓰기</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="small-meta">'
    '버튼 번호가 실행 순서입니다. <b>5-0</b> 자료 스캔(선택) → '
    '<b>5-1~5-4</b> 본문 생성 → <b>점검 A·B·C</b> 는 생성 후 필요할 때만 씁니다.'
    '</div>',
    unsafe_allow_html=True,
)

# ─────────────────────────────────────
# v3.13 M15-B: 계량 수치 점검 (집필 직전 · Unit 설계 재실행 불필요)
# ─────────────────────────────────────
if "metric_scan_result" not in st.session_state:
    st.session_state["metric_scan_result"] = None

_mw_items = st.session_state.get("metric_scan_result")
_mw_count = len(_mw_items) if _mw_items else 0
_mw_label = (
    f"🔍 계량 수치 점검 (M15) — 검출 {_mw_count}건"
    if _mw_items is not None else "🔍 계량 수치 점검 (M15)"
)

with st.expander(_mw_label, expanded=False):
    st.markdown(
        '<div class="callout">'
        '<b>왜 필요한가</b> — 컨셉 카드나 설계안에 "0.5초 간격", "3초간 침묵", "90도" 같은 '
        '계량 수치가 있으면 집필 단계가 그 숫자를 본문에 그대로 옮겨 적습니다. '
        '자료를 다시 만들 필요는 없습니다. 검출된 문구를 집필 프롬프트에 지목 주입해서 차단합니다.'
        '</div>',
        unsafe_allow_html=True,
    )

    st.checkbox(
        "워치리스트 자동 주입 (권장) — 검출된 문구를 집필 프롬프트에서 지목 차단",
        key="metric_watchlist_on",
    )

    if st.button("5-0 · 자료 스캔 (집필 전 점검)", use_container_width=True, key="metric_scan_btn"):
        st.session_state["metric_scan_result"] = scan_metric_expressions()
        st.rerun()

    _items = st.session_state.get("metric_scan_result")
    if _items is None:
        st.caption("‘자료 스캔’을 누르면 컨셉 카드·보강본·Unit 설계·Creator 카드를 검사합니다.")
    elif not _items:
        st.success("계량 수치가 검출되지 않았습니다. 그대로 집필하셔도 됩니다.")
    else:
        _by_where = {}
        for it in _items:
            _by_where.setdefault(it["where"], []).append(it)
        st.markdown(f"**검출 {len(_items)}건 / {len(_by_where)}개 위치**")
        for _where, _group in _by_where.items():
            _exprs = ", ".join(dict.fromkeys(g["expr"] for g in _group))
            st.markdown(f"- **{_where}** — {_exprs}")

        st.markdown("---")
        st.markdown("**자료 직접 수정 (선택)**")
        st.caption(
            "워치리스트만으로도 차단됩니다. 자료 자체를 정리하고 싶을 때만 사용하세요. "
            "자동 치환은 하지 않습니다. 작가가 입력한 문장으로만 한 문장씩 교체합니다."
        )

        _editable = [it for it in _items if it.get("editable")]
        if not _editable:
            st.caption("편집 가능한 항목이 없습니다. (Creator 카드 원본은 읽기 전용입니다.)")
        else:
            _labels = [
                f"[{it['where']}] {it['expr']} — {it['sentence'][:40]}…"
                for it in _editable
            ]
            _sel = st.selectbox(
                "수정할 항목", list(range(len(_editable))),
                format_func=lambda i: _labels[i], key="metric_fix_sel",
            )
            _target = _editable[_sel]
            st.text_area(
                "원문 (읽기 전용)", value=_target["sentence"], height=110,
                disabled=True, key="metric_fix_orig",
            )
            _new = st.text_area(
                "수정 문장 — 숫자를 지우고 몸의 단위·비교·결과로 바꿔 주세요",
                value=_target["sentence"], height=110, key="metric_fix_new",
            )
            _fix_c1, _fix_c2 = st.columns([1, 3])
            with _fix_c1:
                if st.button("이 문장 교체", type="primary", key="metric_fix_apply"):
                    ok = replace_in_source(
                        _target["src_key"], _target["sentence"], _new.strip()
                    )
                    if ok:
                        st.session_state["metric_scan_result"] = scan_metric_expressions()
                        st.success("교체했습니다. 재스캔 완료.")
                        st.rerun()
                    else:
                        st.warning("교체하지 못했습니다. 원문이 이미 변경되었을 수 있습니다. 다시 스캔해 주세요.")
            with _fix_c2:
                st.caption("교체는 해당 자료 필드의 그 문장 1회만 바뀝니다. Unit 본문은 건드리지 않습니다.")

# 집필 프롬프트에 주입할 워치리스트 블록
metric_watchlist_block = ""
if st.session_state.get("metric_watchlist_on") and st.session_state.get("metric_scan_result"):
    metric_watchlist_block = build_metric_watchlist_block(
        st.session_state["metric_scan_result"]
    )


unit_options = [f"{i:02d}" for i in range(1, 13)] + ["13"]
selected_unit = st.selectbox(
    "작업할 Unit 선택",
    unit_options,
    format_func=lambda x: "UNIT 13 · 에필로그" if x == "13" else f"UNIT {x}",
)

# ─── Chapter 1 다단계 생성 시스템 ───
if selected_unit == "01":
    st.markdown(
        '<div class="callout" style="border-left-color:var(--y)">'
        '<b>Chapter 1 다단계 생성</b> — 오프닝은 소설의 얼굴입니다. 3단계로 나눠서 각 단계를 확인하고 승인한 뒤 다음 단계로 넘어갑니다.'
        '<br>Stage A: PEAK (정상) → Stage B: WORLD (전개) → Stage C: LOSS (균열)'
        '</div>',
        unsafe_allow_html=True,
    )

    # v3.14 M16 — 음식 오프닝 시그니처
    _food_c1, _food_c2 = st.columns([2, 3])
    with _food_c1:
        st.checkbox(
            "🍳 음식 오프닝 시그니처 (M16)",
            key="signature_food_opening",
            help=(
                "Chapter 1을 음식·요리·먹는 행위로 시작합니다. "
                "음식 앵커링 4경로(직업/계급/관계/결핍) 중 최소 1개를 통과시켜 "
                "그 인물만의 음식이 되게 합니다. 조리가 불가능한 상황이면 "
                "마시기·씹기로 축소하거나 직업 도구로 폴백합니다."
            ),
        )
    with _food_c2:
        if st.session_state.get("signature_food_opening", True):
            st.caption("작가 시그니처 적용 중 — Stage A와 Unit 01-02 설계에 주입됩니다.")
        else:
            st.caption("시그니처 해제 — 오프닝 소재를 엔진이 자유롭게 선택합니다.")

    # ── v3.16.3 · 1화 전체 재생성 ──
    # 1화는 Stage A/B/C 3단계라 처음부터 다시 하려면 버튼을 세 번 눌러야 했고,
    # 이전 Stage 텍스트가 남아 있으면 새 Stage가 그것을 이어받아 섞였다.
    # 이 버튼은 1화 관련 상태를 한 번에 비워 깨끗한 재시작을 만든다.
    with st.expander("♻️ 1화 처음부터 다시 쓰기", expanded=False):
        st.caption(
            "Stage A·B·C 원고와 UNIT 01 확정본, 회차 제목, 상태 원장, 요약을 모두 비웁니다. "
            "설계안(STEP 4)과 STEP 1 자료는 그대로 둡니다."
        )
        _r1c1, _r1c2 = st.columns([1, 1])
        with _r1c1:
            _reset1_ok = st.checkbox("네, 1화를 비우겠습니다", key="ch1_reset_confirm")
        with _r1c2:
            if st.button(
                "1화 초기화 후 다시 시작",
                use_container_width=True,
                disabled=not _reset1_ok,
                key="ch1_reset_btn",
            ):
                for _k in ("ch1_stage_a", "ch1_stage_b", "ch1_stage_c"):
                    st.session_state[_k] = ""
                st.session_state["unit_drafts"]["01"] = ""
                st.session_state["chapter_titles"]["01"] = ""
                for _d in ("unit_summaries", "continuity_ledger", "character_tracker"):
                    if isinstance(st.session_state.get(_d), dict):
                        st.session_state[_d].pop("01", None)
                st.session_state["quality_report"] = {}
                st.session_state["ch1_reset_confirm"] = False
                set_status("1화를 비웠습니다. 5-1 Stage A부터 다시 시작하세요.", "success")
                st.rerun()

    # v3.15.2 — 진행 표시. 버튼이 4개라 어디까지 했는지 한눈에 안 보였다.
    _done_a = bool(st.session_state.get("ch1_stage_a", "").strip())
    _done_b = bool(st.session_state.get("ch1_stage_b", "").strip())
    _done_c = bool(st.session_state.get("ch1_stage_c", "").strip())
    _done_m = bool(st.session_state["unit_drafts"].get("01", "").strip())
    _steps = [("5-1 Stage A", _done_a), ("5-2 Stage B", _done_b),
              ("5-3 Stage C", _done_c), ("5-4 확정", _done_m)]
    st.markdown(
        '<div class="small-meta">진행 — '
        + "  →  ".join(f"{'✅' if d else '⬜'} {lbl}" for lbl, d in _steps)
        + "</div>",
        unsafe_allow_html=True,
    )

    ch1_a, ch1_b, ch1_c = st.columns(3)

    with ch1_a:
        st.markdown("**Stage A · PEAK**")
        _sigtag = "음식 시그니처" if st.session_state.get("signature_food_opening", True) else "소재 자유"
        st.markdown(f'<div class="small-meta">오프닝 장면 · {_sigtag} · 인물 정의 · ~2000자</div>', unsafe_allow_html=True)
        if st.button("5-1 · Stage A 생성", type="primary", use_container_width=True, key="ch1_a_btn"):
            def _job():
                prompt = build_ch1_stage_a_prompt(
                    working_title=working_title, genre=genre, format_mode=format_mode,
                    pov=pov, overview=overview, characters=characters,
                    synopsis=synopsis, notes=notes,
                    style_dna=st.session_state["style_dna"], style_strength=style_strength,
                    locked_block=locked_block,
                    # v3.0 신규
                    profession_text=profession_text_combined,
                    period_keys=active_period_keys,
                    # v3.13 M15-B
                    metric_watchlist=metric_watchlist_block,
                    # v3.14 M16
                    food_signature=st.session_state.get("signature_food_opening", True),
                    # v3.15 M19 — UNIT 01 설계안 주입 (v3.14까지 미전달이던 경로)
                    all_blueprints_text=all_blueprints_text,
                )
                return llm_call(prompt, max_tokens=MAX_TOKENS_LONG, use_opus=True)
            result = run_with_status("Stage A: PEAK 오프닝을 생성 중입니다...", "Stage A 생성 완료.", _job)
            if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
                # v3.15 — Stage 단위에서 '끝.'·마크다운 잔재를 즉시 걷어낸다
                result, _slog = sanitize_manuscript(result, is_final_unit=False)
                if _slog:
                    st.caption("본문 정리: " + " / ".join(_slog))
                st.session_state["ch1_stage_a"] = result

    with ch1_b:
        st.markdown("**Stage B · WORLD**")
        st.markdown('<div class="small-meta">세계관 · 관계 · 권력 구조를 장면 안에 · ~2500자</div>', unsafe_allow_html=True)
        has_a = bool(st.session_state["ch1_stage_a"].strip())
        if st.button("5-2 · Stage B 생성", use_container_width=True, disabled=not has_a, key="ch1_b_btn"):
            def _job():
                prompt = build_ch1_stage_b_prompt(
                    working_title=working_title, genre=genre, format_mode=format_mode,
                    pov=pov, overview=overview, characters=characters,
                    synopsis=synopsis, notes=notes,
                    style_dna=st.session_state["style_dna"], style_strength=style_strength,
                    stage_a_text=st.session_state["ch1_stage_a"],
                    locked_block=locked_block,
                    # v3.0 신규
                    profession_text=profession_text_combined,
                    period_keys=active_period_keys,
                    # v3.13 M15-B
                    metric_watchlist=metric_watchlist_block,
                    # v3.15 M19
                    all_blueprints_text=all_blueprints_text,
                )
                return llm_call(prompt, max_tokens=MAX_TOKENS_LONG, use_opus=True)
            result = run_with_status("Stage B: WORLD 전개를 생성 중입니다...", "Stage B 생성 완료.", _job)
            if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
                # v3.15 — Stage B가 출력한 '끝.'이 병합본 중간에 박히던 사고 차단
                result, _slog = sanitize_manuscript(result, is_final_unit=False)
                if _slog:
                    st.caption("본문 정리: " + " / ".join(_slog))
                st.session_state["ch1_stage_b"] = result

    with ch1_c:
        st.markdown("**Stage C · LOSS**")
        st.markdown('<div class="small-meta">균열 · 상실의 신호 · 클리프행어 · ~1500자</div>', unsafe_allow_html=True)
        has_b = bool(st.session_state["ch1_stage_b"].strip())
        if st.button("5-3 · Stage C 생성", use_container_width=True, disabled=not has_b, key="ch1_c_btn"):
            def _job():
                prompt = build_ch1_stage_c_prompt(
                    working_title=working_title, genre=genre, format_mode=format_mode,
                    pov=pov, overview=overview, characters=characters,
                    synopsis=synopsis, notes=notes,
                    style_dna=st.session_state["style_dna"], style_strength=style_strength,
                    stage_a_text=st.session_state["ch1_stage_a"],
                    stage_b_text=st.session_state["ch1_stage_b"],
                    locked_block=locked_block,
                    # v3.0 신규
                    profession_text=profession_text_combined,
                    period_keys=active_period_keys,
                    # v3.13 M15-B
                    metric_watchlist=metric_watchlist_block,
                    # v3.15 M19
                    all_blueprints_text=all_blueprints_text,
                )
                return llm_call(prompt, max_tokens=MAX_TOKENS_LONG, use_opus=True)
            result = run_with_status("Stage C: LOSS 균열을 생성 중입니다...", "Stage C 생성 완료.", _job)
            if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
                result, _slog = sanitize_manuscript(result, is_final_unit=False)
                if _slog:
                    st.caption("본문 정리: " + " / ".join(_slog))
                st.session_state["ch1_stage_c"] = result

    # 각 Stage 미리보기
    for stage_key, stage_label in [("ch1_stage_a", "Stage A · PEAK"), ("ch1_stage_b", "Stage B · WORLD"), ("ch1_stage_c", "Stage C · LOSS")]:
        stage_text = st.session_state.get(stage_key, "")
        if stage_text.strip():
            with st.expander(f"{stage_label} 보기", expanded=True):
                st.text_area(stage_label, value=stage_text, height=300, label_visibility="collapsed", key=f"preview_{stage_key}")

    # 3단계 완성 시 합치기 버튼
    all_stages_done = all(st.session_state.get(k, "").strip() for k in ["ch1_stage_a", "ch1_stage_b", "ch1_stage_c"])
    if all_stages_done:
        st.markdown("---")
        if st.button("5-4 · Chapter 1 확정 — 3단계를 합쳐서 UNIT 01로 저장", type="primary", use_container_width=True, key="ch1_merge_btn"):
            merged = (
                st.session_state["ch1_stage_a"].strip()
                + "\n\n"
                + st.session_state["ch1_stage_b"].strip()
                + "\n\n"
                + st.session_state["ch1_stage_c"].strip()
            )
            # v3.15 — 병합 직후 잔여 마커 정리 (2차 방어)
            # Stage 단위에서 이미 걷어내지만, 옛 세션에서 이어받은 Stage 텍스트나
            # 작가가 직접 붙여넣은 텍스트에도 '끝.'이 남아 있을 수 있다.
            merged, merge_log = sanitize_manuscript(merged, is_final_unit=False)
            if merge_log:
                st.info("병합 시 정리한 항목 — " + " / ".join(merge_log))

            # 챕터 제목 파싱 — v3.16: 설계안 제목이 정본
            ch_title, ch_body = parse_chapter_title(merged)
            st.session_state["unit_drafts"]["01"] = ch_body if ch_title else merged
            _final_title, _title_note = resolve_chapter_title(1, ch_title, all_blueprints_text)
            st.session_state["chapter_titles"]["01"] = _final_title
            st.caption(f"회차 제목 — {_final_title}")
            if _title_note:
                st.info(_title_note)
            set_status("Chapter 1이 확정되었습니다. UNIT 01로 저장 완료.", "success")
            # 품질 자동 체크
            final_text = ch_body if ch_title else merged
            qr = analyze_unit_quality(final_text)
            st.session_state["quality_report"] = qr

            # v3.15 — Chapter 1 경로에도 임계치 초과 경고를 붙인다.
            # v3.14까지는 리포트만 조용히 저장해서, '있었다' 16회(임계치 10)가
            # 그대로 통과했다. Stage 단위 재생성이 필요하므로 자동 재생성은
            # 하지 않고, 어느 Stage를 다시 뽑아야 하는지 작가에게 알린다.
            if qr.get("should_regenerate"):
                _hard = [
                    f"{k} {v.get('count')}회 (임계치 {v.get('threshold')})"
                    for k, v in qr.get("violations", {}).items()
                    if v.get("severity") in ("critical", "high")
                ]
                st.error(
                    "⚠️ **BJND 임계치 초과 상태로 저장되었습니다.** — "
                    + " / ".join(_hard)
                    + "  \n원고는 저장됐으니 그대로 두셔도 되고, Stage를 다시 뽑아 "
                    "교체하셔도 됩니다. 어느 Stage에서 몰렸는지는 아래 Stage 보기에서 "
                    "확인할 수 있습니다."
                )

            # Unit 요약 자동 생성
            summary = generate_unit_summary(1, final_text)
            if summary:
                if "unit_summaries" not in st.session_state:
                    st.session_state["unit_summaries"] = {}
                st.session_state["unit_summaries"]["01"] = summary
            # v3.15 M17 — 연속성 상태 원장 생성
            generate_continuity_ledger(1, final_text)
            # 캐릭터 등장 추적
            track_characters("01", final_text)

# ─── 일반 Unit 생성 (Unit 02~13) ───
else:
    draft_col1, draft_col2, draft_col3 = st.columns([1, 1, 1])

    with draft_col1:
        if st.button("5-1 · Unit 원고 생성", type="primary", use_container_width=True):
            unit_no = int(selected_unit)

            if unit_no == 13:
                def _job():
                    prompt = build_epilogue_prompt(
                        working_title=working_title,
                        genre=genre,
                        overview=overview,
                        characters=characters,
                        synopsis=synopsis,
                        story_reinforcement_merged=story_merged_text,
                        all_blueprints_text=all_blueprints_text,
                        all_drafts_text=gather_recent_drafts(13, window=3),
                        style_dna=st.session_state["style_dna"],
                        locked_block=locked_block,
                    )
                    return generate_or_expand_unit(13, prompt)

                result = run_with_status(
                    "UNIT 13 에필로그를 생성 중입니다...",
                    "UNIT 13 에필로그 생성이 완료되었습니다.",
                    _job,
                )
                if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
                    # v3.15 — 에필로그는 마지막 '끝.'을 보존한다.
                    result, _slog = sanitize_manuscript(result, is_final_unit=True)
                    if _slog:
                        st.caption("본문 정리: " + " / ".join(_slog))
                    ch_title, ch_body = parse_chapter_title(result)
                    st.session_state["unit_drafts"][selected_unit] = ch_body if ch_title else result
                    # v3.16 — 에필로그는 설계 그룹이 없으므로 고정 라벨
                    st.session_state["chapter_titles"][selected_unit] = (
                        ch_title if (ch_title and "—" in ch_title) else "[CHAPTER 13] — 에필로그"
                    )
            else:
                def _job():
                    # v3.0 M1: 자동 재생성 로직
                    # 1차 생성
                    prompt = build_unit_draft_prompt(
                        unit_no=unit_no,
                        working_title=working_title,
                        genre=genre,
                        format_mode=format_mode,
                        pov=pov,
                        overview=overview,
                        characters=characters,
                        synopsis=synopsis,
                        notes=notes,
                        story_reinforcement_merged=story_merged_text,
                        all_blueprints_text=all_blueprints_text,
                        previous_drafts=gather_recent_drafts(unit_no),
                        style_dna=st.session_state["style_dna"],
                        style_strength=style_strength,
                        target_length=UNIT_TARGET_LENGTHS.get(unit_no, 8000),
                        min_length=UNIT_MIN_LENGTHS.get(unit_no, 6000),
                        locked_block=locked_block,
                        # v3.0 신규
                        profession_text=profession_text_combined,
                        period_keys=active_period_keys,
                        metric_watchlist=metric_watchlist_block,
                        food_signature=st.session_state.get("signature_food_opening", True),
                        retry_hint="",
                        # v3.15 M17 — 직전까지의 물리 상태 원장 주입
                        continuity_ledger=gather_continuity_ledger(unit_no),
                    )
                    first_result = generate_or_expand_unit(unit_no, prompt)

                    # v3.0 M1: 1차 결과 BJND 검증
                    if first_result:
                        check_body = parse_chapter_title(first_result)[1] or first_result
                        qr_first = analyze_unit_quality(check_body)

                        # 재생성 트리거 조건 확인
                        if qr_first.get("should_regenerate") and AUTO_REGEN_MAX_RETRIES > 0:
                            st.info(
                                f"⚡ BJND Scene Enforcer 발동: 임계치 초과 감지. "
                                f"자동 재생성 시도 중... "
                                f"(위반: {', '.join(qr_first.get('violations', {}).keys())})"
                            )
                            # 위반 지표를 힌트로 구성
                            retry_hint = build_retry_hint(qr_first.get("violations", {}))
                            # 2차 재생성 (retry_hint 주입)
                            retry_prompt = build_unit_draft_prompt(
                                unit_no=unit_no,
                                working_title=working_title,
                                genre=genre,
                                format_mode=format_mode,
                                pov=pov,
                                overview=overview,
                                characters=characters,
                                synopsis=synopsis,
                                notes=notes,
                                story_reinforcement_merged=story_merged_text,
                                all_blueprints_text=all_blueprints_text,
                                previous_drafts=gather_recent_drafts(unit_no),
                                style_dna=st.session_state["style_dna"],
                                style_strength=style_strength,
                                target_length=UNIT_TARGET_LENGTHS.get(unit_no, 8000),
                                min_length=UNIT_MIN_LENGTHS.get(unit_no, 6000),
                                locked_block=locked_block,
                                profession_text=profession_text_combined,
                                period_keys=active_period_keys,
                                metric_watchlist=metric_watchlist_block,
                                food_signature=st.session_state.get("signature_food_opening", True),
                                retry_hint=retry_hint,
                                # v3.15 M17
                                continuity_ledger=gather_continuity_ledger(unit_no),
                            )
                            retry_result = generate_or_expand_unit(unit_no, retry_prompt)

                            if retry_result:
                                # 2차 결과가 더 좋으면 교체, 아니면 1차 유지
                                retry_body = parse_chapter_title(retry_result)[1] or retry_result
                                qr_retry = analyze_unit_quality(retry_body)
                                # severity high 이상 위반 수 비교
                                def count_serious(qr):
                                    return sum(
                                        1 for v in qr.get("violations", {}).values()
                                        if v.get("severity") in ("critical", "high")
                                    )
                                if count_serious(qr_retry) < count_serious(qr_first):
                                    st.success("✅ 재생성본이 더 좋습니다. 재생성본으로 교체.")
                                    return retry_result
                                else:
                                    st.warning("⚠️ 재생성본이 개선되지 않음. 1차 생성본 유지.")
                    return first_result

                done_msg = f"UNIT {unit_no:02d} 원고 생성이 완료되었습니다."
                if unit_no == 12:
                    done_msg = "UNIT 12 본편 마무리 생성이 완료되었습니다."

                result = run_with_status(
                    f"UNIT {unit_no:02d} 원고를 생성 중입니다...",
                    done_msg,
                    _job,
                )
                if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
                    # v3.15 — 잔여 마커 정리. Unit 12는 마지막 '끝.'을 보존한다.
                    result, _slog = sanitize_manuscript(result, is_final_unit=(unit_no in (12, 13)))
                    if _slog:
                        st.caption("본문 정리: " + " / ".join(_slog))
                    # v3.16 — 제목은 설계안이 정본. 집필 단계 서브타이틀은 폐지.
                    ch_title, ch_body = parse_chapter_title(result)
                    st.session_state["unit_drafts"][selected_unit] = ch_body if ch_title else result
                    _final_title, _title_note = resolve_chapter_title(
                        unit_no, ch_title, all_blueprints_text
                    )
                    st.session_state["chapter_titles"][selected_unit] = _final_title
                    st.caption(f"회차 제목 — {_final_title}")
                    if _title_note:
                        st.info(_title_note)
                    check_text = ch_body if ch_title else result
                    if is_incomplete_text(check_text, unit_no):
                        set_status(
                            f"UNIT {unit_no:02d}는 생성되었지만 아직 짧거나 미완성일 수 있습니다. 다시 쓰기나 재생성을 권장합니다.",
                            "warning",
                        )
                    # 품질 자동 체크
                    qr = analyze_unit_quality(check_text)
                    # v3.15.1 — 앞 Unit들과의 교차 반복 진단을 합류시킨다
                    _cross = analyze_cross_unit_repetition(unit_no, check_text)
                    if _cross:
                        qr["issues"] = list(qr.get("issues", [])) + _cross
                    st.session_state["quality_report"] = qr
                    # Unit 요약 자동 생성
                    summary = generate_unit_summary(unit_no, check_text)
                    if summary:
                        if "unit_summaries" not in st.session_state:
                            st.session_state["unit_summaries"] = {}
                        st.session_state["unit_summaries"][selected_unit] = summary
                    # v3.15 M17 — 연속성 상태 원장 생성
                    generate_continuity_ledger(unit_no, check_text)
                    # 캐릭터 등장 추적
                    track_characters(selected_unit, check_text)

    with draft_col2:
        # ── v3.16.3 · 이 Unit 처음부터 다시 쓰기 ──
        # '다시 쓰기'는 기존 원고를 손보는 것이고, 이것은 완전히 새로 뽑는 것이다.
        # 원고를 비워야 previous_drafts에 옛 원고가 섞여 들어가지 않는다.
        with st.expander(f"♻️ {selected_unit}화 처음부터 다시 쓰기", expanded=False):
            st.caption(
                "이 Unit의 원고·제목·상태 원장·요약을 비웁니다. "
                "비운 뒤 5-1로 새로 생성하세요. 설계안은 그대로 둡니다."
            )
            _ru_ok = st.checkbox(
                f"네, {selected_unit}화를 비우겠습니다", key=f"unit_reset_confirm_{selected_unit}"
            )
            if st.button(
                f"{selected_unit}화 초기화",
                use_container_width=True,
                disabled=not _ru_ok,
                key=f"unit_reset_btn_{selected_unit}",
            ):
                st.session_state["unit_drafts"][selected_unit] = ""
                st.session_state["chapter_titles"][selected_unit] = ""
                for _d in ("unit_summaries", "continuity_ledger", "character_tracker"):
                    if isinstance(st.session_state.get(_d), dict):
                        st.session_state[_d].pop(selected_unit, None)
                st.session_state["quality_report"] = {}
                st.session_state[f"unit_reset_confirm_{selected_unit}"] = False
                set_status(
                    f"{selected_unit}화를 비웠습니다. 5-1로 새로 생성하세요.", "success"
                )
                st.rerun()

        rewrite_mode = st.selectbox(
            "다시 쓰기 모드",
            ["더 상업적으로", "더 빠르게", "더 감정적으로", "더 차갑게", "더 영상적으로", "더 문학적으로"],
            index=0,
        )
        if st.button("5-2 · Unit 다시 쓰기 (선택)", use_container_width=True):
            source_text = st.session_state["unit_drafts"].get(selected_unit, "")
            if source_text.strip():
                def _job():
                    prompt = build_unit_rewrite_prompt(
                        unit_no=int(selected_unit),
                        rewrite_mode=rewrite_mode,
                        source_text=source_text,
                        style_dna=st.session_state["style_dna"],
                        target_length=UNIT_TARGET_LENGTHS.get(int(selected_unit), 8000),
                        min_length=UNIT_MIN_LENGTHS.get(int(selected_unit), 6000),
                    )
                    return generate_or_expand_unit(int(selected_unit), prompt)

                result = run_with_status(
                    f"UNIT {selected_unit}를 다시 쓰는 중입니다...",
                    f"UNIT {selected_unit} 다시 쓰기가 완료되었습니다.",
                    _job,
                )
                if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
                    st.session_state["unit_drafts"][selected_unit] = result

    with draft_col3:
        unit_num = int(selected_unit)
        st.markdown(
            f'<div class="small-meta">목표 분량 {UNIT_TARGET_LENGTHS.get(unit_num, 8000):,}자 / 최소 {UNIT_MIN_LENGTHS.get(unit_num, 6000):,}자</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            '<div class="small-meta">UNIT 12는 본편을 반드시 마무리합니다. UNIT 13은 선택형 에필로그입니다.</div>',
            unsafe_allow_html=True,
        )

current_draft = st.session_state["unit_drafts"].get(selected_unit, "")
current_ch_title = st.session_state["chapter_titles"].get(selected_unit, "")
if current_draft:
    label = current_ch_title if current_ch_title else (
        "UNIT 13 · 에필로그" if selected_unit == "13" else f"UNIT {selected_unit}"
    )
    with st.expander(f"{label} 보기", expanded=True):
        st.text_area("원고", value=current_draft, height=420, label_visibility="collapsed")

    # ── v3.15 신규 도구 3종 ──
    # 작업 규칙 7 준수 — 본문 문장은 자동으로 고치지 않는다.
    # 마커 제거·재검사·원장 생성만 제공하고, 문장 수정은 작가가 직접 한다.
    tool_c1, tool_c2, tool_c3 = st.columns(3)

    with tool_c1:
        if st.button("점검 A · 이 Unit 재검사", use_container_width=True, key="requalify_btn"):
            _qr = analyze_unit_quality(current_draft)
            _cross = analyze_cross_unit_repetition(int(selected_unit), current_draft)
            if _cross:
                _qr["issues"] = list(_qr.get("issues", [])) + _cross
            st.session_state["quality_report"] = _qr
            set_status(f"UNIT {selected_unit} 재검사 완료.", "success")

    with tool_c2:
        # 정리할 것이 있는지 미리 계산해 버튼 상태를 결정한다.
        _preview, _plog = sanitize_manuscript(
            current_draft, is_final_unit=(selected_unit in ("12", "13"))
        )
        if st.button(
            "점검 B · 본문 정리",
            use_container_width=True,
            disabled=not _plog,
            key="sanitize_btn",
            help="본문 중간 '끝.'과 마크다운 기호만 제거합니다. 문장은 건드리지 않습니다.",
        ):
            st.session_state["unit_drafts"][selected_unit] = _preview
            st.session_state["quality_report"] = analyze_unit_quality(_preview)
            set_status("본문 정리 완료 — " + " / ".join(_plog), "success")
            st.rerun()
        if _plog:
            st.caption("정리 대상: " + " / ".join(_plog))
        else:
            st.caption("정리할 마커 없음")

    with tool_c3:
        if st.button(
            "점검 C · 상태 원장 생성",
            use_container_width=True,
            key="ledger_btn",
            help="다음 Unit 집필 시 착의·소지품·시각·날씨가 어긋나지 않도록 원장을 만듭니다. (M17)",
        ):
            with st.spinner("연속성 상태 원장을 추출 중입니다..."):
                _led = generate_continuity_ledger(int(selected_unit), current_draft)
            if _led:
                set_status(f"UNIT {selected_unit} 상태 원장 생성 완료.", "success")
            else:
                set_status("상태 원장 생성에 실패했습니다. API 키와 연결을 확인하세요.", "error")

    # 저장된 원장 표시
    _led_saved = (st.session_state.get("continuity_ledger", {}) or {}).get(
        f"{int(selected_unit):02d}", ""
    )
    if _led_saved:
        with st.expander(f"📋 UNIT {selected_unit} 연속성 상태 원장 (M17)", expanded=False):
            st.caption(
                "다음 Unit 집필 프롬프트에 그대로 주입됩니다. "
                "사실이 틀렸으면 아래에서 직접 고쳐주세요."
            )
            _led_edit = st.text_area(
                "상태 원장",
                value=_led_saved,
                height=280,
                label_visibility="collapsed",
                key=f"ledger_edit_{selected_unit}",
            )
            if _led_edit != _led_saved:
                if st.button("원장 수정 저장", key=f"ledger_save_{selected_unit}"):
                    st.session_state["continuity_ledger"][f"{int(selected_unit):02d}"] = _led_edit
                    set_status(f"UNIT {selected_unit} 원장을 수정했습니다.", "success")
                    st.rerun()

# 품질 리포트 표시
qr = st.session_state.get("quality_report", {})
if qr.get("issues") or qr.get("stats"):
    with st.expander("📊 품질 리포트", expanded=True):
        stats = qr.get("stats", {})
        if stats:
            stat_cols = st.columns(4)
            stat_items = list(stats.items())
            for idx, (k, v) in enumerate(stat_items[:4]):
                with stat_cols[idx]:
                    st.metric(k, v)
            if len(stat_items) > 4:
                stat_cols2 = st.columns(4)
                for idx, (k, v) in enumerate(stat_items[4:8]):
                    with stat_cols2[idx]:
                        st.metric(k, v)
        issues = qr.get("issues", [])
        if issues:
            for issue in issues:
                st.warning(issue)
        else:
            st.success("✅ 주요 품질 문제 없음")

# Unit 요약 표시
summaries = st.session_state.get("unit_summaries", {})
filled = {k: v for k, v in summaries.items() if v}
if filled:
    with st.expander("📋 Unit 요약 (전체 흐름)", expanded=False):
        for key in sorted(filled.keys()):
            st.markdown(f"**UNIT {key}**: {filled[key]}")

# 캐릭터 등장 추적 표시
char_report = get_character_report()
if char_report.get("first_appearance"):
    with st.expander(f"👥 캐릭터 등장 추적 ({char_report.get('total', 0)}명)", expanded=False):
        fa = char_report["first_appearance"]
        for name in sorted(fa.keys(), key=lambda x: fa[x]):
            st.markdown(f"- **{name}** — 첫 등장: UNIT {fa[name]}")
        warnings = char_report.get("warnings", [])
        if warnings:
            st.markdown("---")
            for w in warnings:
                st.warning(w)

# ─────────────────────────────────────
# STEP 6
# ─────────────────────────────────────
st.markdown('<div class="section-header">💾 STEP 6 · 저장 / 내보내기</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="small-meta">현재 Unit만 저장하거나, 전체 통합본을 저장합니다. '
    '작업을 이어서 하려면 상단의 프로젝트 저장(JSON)을 함께 쓰세요.</div>',
    unsafe_allow_html=True,
)

safe_title = safe_filename(working_title)
manuscript = final_manuscript_text(working_title)

current_unit_text = st.session_state["unit_drafts"].get(selected_unit, "").strip()
current_unit_label = "UNIT_13_에필로그" if selected_unit == "13" else f"UNIT_{selected_unit}"

# ── DOCX 문단 간격 (v3.16.2) ──
# 기존에는 본문의 빈 줄을 전부 빈 단락으로 옮겨 원고 분량이 두 배가 됐다.
# (실측 — UNIT 01: 219단락 중 108개가 빈 단락, 49%)
_spacing_labels = {
    "standard": "표준 — 들여쓰기로 문단 구분 (한국 소설 원고 기본)",
    "relaxed": "여유 — 들여쓰기 + 문단 뒤 6pt 여백",
    "web": "웹 연재 — 빈 줄을 그대로 유지",
}
_spacing_mode = st.radio(
    "DOCX 문단 간격",
    options=list(_spacing_labels.keys()),
    format_func=lambda k: _spacing_labels[k],
    horizontal=True,
    index=0,
    key="docx_spacing_mode",
    help="빈 줄 2개 이상이나 장면 전환 마커는 어느 모드에서도 여백으로 남습니다.",
)

txt_bytes = export_txt(manuscript) if manuscript.strip() else b""
docx_bytes = (
    export_docx(working_title or "Novel Draft", manuscript, _spacing_mode)
    if manuscript.strip() else b""
)

unit_txt_bytes = export_txt(current_unit_text) if current_unit_text else b""
unit_docx_bytes = (
    export_docx(
        f"{working_title or 'Novel Draft'} {current_unit_label}",
        current_unit_text,
        _spacing_mode,
    )
    if current_unit_text
    else b""
)

if not manuscript.strip():
    st.warning("아직 저장할 최종 원고가 없습니다. 먼저 Unit 원고를 생성해 주세요.")
else:
    st.info("다운로드 버튼을 누르면 브라우저로 바로 저장됩니다.")

st.markdown("**현재 Unit 저장**")
u1, u2 = st.columns(2)

with u1:
    st.download_button(
        "현재 Unit TXT 저장",
        data=unit_txt_bytes,
        file_name=f"{safe_title}_{current_unit_label}.txt",
        mime="text/plain",
        use_container_width=True,
        disabled=not bool(current_unit_text),
        key=f"download_unit_txt_{selected_unit}",
    )

with u2:
    st.download_button(
        "현재 Unit DOCX 저장",
        data=unit_docx_bytes,
        file_name=f"{safe_title}_{current_unit_label}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        disabled=not bool(current_unit_text),
        key=f"download_unit_docx_{selected_unit}",
    )

st.markdown("**최종 원고 저장**")
exp1, exp2 = st.columns(2)

with exp1:
    st.download_button(
        "최종 원고 TXT 저장",
        data=txt_bytes,
        file_name=f"{safe_title}_final.txt",
        mime="text/plain",
        use_container_width=True,
        disabled=not bool(manuscript.strip()),
        key="download_final_txt",
    )

with exp2:
    st.download_button(
        "최종 원고 DOCX 저장",
        data=docx_bytes,
        file_name=f"{safe_title}_final.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        disabled=not bool(manuscript.strip()),
        key="download_final_docx",
    )

with st.expander("최종 원고 미리보기", expanded=False):
    st.text_area("최종 원고", value=manuscript, height=420, label_visibility="collapsed")


# ─────────────────────────────────────
# STEP 7
# v3.15.2 — 기존에는 이 블록이 STEP 5(원고 생성)와 STEP 7(저장) 사이에
# STEP 6으로 끼어 있었다. 작품 전체 제목은 완고 후 한 번만 정하는 일인데,
# Unit을 한 편 생성할 때마다 저장하러 내려가는 길목에 놓여 있었다.
# 저장 뒤로 옮기고, 기본 접힘 상태로 둔다.
# ─────────────────────────────────────
_title_done = bool(
    st.session_state.get("title_review", "").strip()
    or st.session_state.get("chapter_title_review", "").strip()
)
_units_written = count_written_units()

with st.expander(
    "🏷️ STEP 7 · 제목 검수 (작품 제목 · 회차 제목)"
    + ("  —  검토 완료" if _title_done else "  —  완고 후 1회"),
    expanded=False,
):
    st.caption(
        "완고 후 한 번만 하는 단계입니다. 회차마다 내려올 필요가 없습니다."
        + (f"  현재 Unit {_units_written}개 작성됨." if _units_written else "")
    )

    _t_tab1, _t_tab2 = st.tabs(["작품 제목", "회차 제목 검수"])

    # ── 7-1 작품 전체 제목 ──
    with _t_tab1:
        title_col1, title_col2 = st.columns([1, 1])
        with title_col1:
            if st.button("7-1 · 원고 기반 작품 제목 검토", use_container_width=True):
                def _job():
                    prompt = build_title_review_prompt(
                        current_title=working_title,
                        overview=overview,
                        synopsis=synopsis,
                        story_reinforcement_merged=story_merged_text,
                        all_blueprints_text=all_blueprints_text,
                        all_drafts_text=gather_all_drafts_text(),
                        style_dna=st.session_state["style_dna"],
                    )
                    return llm_call(prompt, max_tokens=MAX_TOKENS_SHORT)

                result = run_with_status(
                    "원고를 다시 읽고 작품 제목을 검토 중입니다...",
                    "작품 제목 검토가 완료되었습니다.",
                    _job,
                )
                if result:  # v3.4.1 — 빈 문자열도 저장하지 않는다
                    st.session_state["title_review"] = result

        with title_col2:
            st.markdown(
                '<div class="small-meta">가제를 버리는 단계가 아니라, 원고를 읽고 현재 가제가 '
                '맞는지 검토하고 대안을 비교하는 단계입니다.</div>',
                unsafe_allow_html=True,
            )

        if st.session_state["title_review"]:
            st.markdown("---")
            st.markdown(st.session_state["title_review"])

    # ── 7-2 회차 제목 검수 (v3.16 신규) ──
    with _t_tab2:
        _cur_titles = gather_chapter_titles_text()
        st.markdown(
            '<div class="small-meta">'
            '회차 제목은 STEP 4 설계에서 전체 맥락을 보고 뽑혀 이미 배열이 잡혀 있습니다. '
            '여기서는 새로 짓지 않고 <b>본문 이탈 · 목차 스포일러 · 이미지 중복</b> '
            '세 가지만 검수합니다. 멀쩡한 제목은 그대로 둡니다.</div>',
            unsafe_allow_html=True,
        )

        if not _cur_titles:
            st.warning("아직 확정된 회차가 없습니다. 먼저 STEP 5에서 Unit 원고를 생성해 주세요.")
        else:
            st.markdown("**현재 목차**")
            st.code(_cur_titles, language=None)

            _ct_c1, _ct_c2 = st.columns([1, 1])
            with _ct_c1:
                if st.button("7-2 · 회차 제목 검수", use_container_width=True, type="primary"):
                    def _job():
                        prompt = build_chapter_title_review_prompt(
                            working_title=working_title,
                            genre=genre,
                            current_titles_text=_cur_titles,
                            all_drafts_text=gather_all_drafts_text(),
                            all_blueprints_text=all_blueprints_text,
                        )
                        return llm_call(prompt, max_tokens=MAX_TOKENS_ANALYSIS)

                    result = run_with_status(
                        "완성 원고와 회차 제목을 대조하는 중입니다...",
                        "회차 제목 검수가 완료되었습니다.",
                        _job,
                    )
                    if result:
                        st.session_state["chapter_title_review"] = result

            with _ct_c2:
                st.markdown(
                    '<div class="small-meta">검수 결과는 제안일 뿐입니다. '
                    '제목 교체는 아래에서 작가가 직접 하세요.</div>',
                    unsafe_allow_html=True,
                )

            if st.session_state.get("chapter_title_review", "").strip():
                st.markdown("---")
                st.markdown(st.session_state["chapter_title_review"])

            # ── 제목 직접 수정 ──
            # ★ 자동 반영하지 않는다. 작가가 고른 것만 들어간다 (작업 규칙 7). ★
            st.markdown("---")
            with st.expander("✏️ 회차 제목 직접 수정", expanded=False):
                st.caption(
                    "검수 제안 중 마음에 드는 것만 골라 직접 넣으세요. "
                    "엔진은 제목을 자동으로 바꾸지 않습니다."
                )
                _edited = {}
                for _i in range(1, 14):
                    _k = f"{_i:02d}" if _i < 13 else "13"
                    if not (st.session_state["unit_drafts"].get(_k) or "").strip():
                        continue
                    _cur = st.session_state["chapter_titles"].get(_k, "")
                    _sub = ""
                    _m = re.match(r"^\[CHAPTER[^\]]*\]\s*[—\-–]\s*(.+)$", _cur.strip())
                    if _m:
                        _sub = _m.group(1).strip()
                    _edited[_k] = st.text_input(
                        f"UNIT {_i:02d} 제목",
                        value=_sub,
                        key=f"ct_edit_{_k}",
                        placeholder="비워두면 번호만 표시됩니다",
                    )
                if st.button("회차 제목 일괄 저장", use_container_width=True, key="ct_save_all"):
                    _changed = 0
                    for _k, _v in _edited.items():
                        _num = int(_k)
                        _new = f"[CHAPTER {_num}] — {_v.strip()}" if _v.strip() else f"[CHAPTER {_num}]"
                        if st.session_state["chapter_titles"].get(_k, "") != _new:
                            st.session_state["chapter_titles"][_k] = _new
                            _changed += 1
                    set_status(f"회차 제목 {_changed}개를 수정했습니다.", "success")
                    st.rerun()

# ─────────────────────────────────────